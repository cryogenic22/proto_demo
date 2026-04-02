"""
TDD tests for document formatter fidelity.

These tests define what "done" looks like for the site contract generator
and DOCX renderer. They must ALL PASS before shipping.
"""

import json
import pytest
from pathlib import Path
from unittest.mock import patch

from src.formatter.extractor import FormattingExtractor, FormattedDocument, FormattedSpan
from src.formatter.docx_renderer import DOCXRenderer
from src.formatter.fidelity_checker import DocumentFidelityChecker, detect_runon_words
from src.formatter.formula_detector import FormulaDetector


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

TEMPLATE_PATH = Path("C:/Users/kapil/Documents/051 ANIMATE CTSA template_15-05-18_22062018_0.pdf")
PROTOCOL_PATH = Path("data_seed/protocols/pfizer_bnt162.json")

def _load_template_if_available():
    """Load template PDF bytes if available, skip test if not."""
    if not TEMPLATE_PATH.exists():
        pytest.skip("Template PDF not available")
    return TEMPLATE_PATH.read_bytes()

def _load_protocol():
    if not PROTOCOL_PATH.exists():
        pytest.skip("Protocol data not available")
    with open(PROTOCOL_PATH) as f:
        return json.load(f)


# ---------------------------------------------------------------------------
# Test 1: FormattingExtractor must capture ALL formatting attributes
# ---------------------------------------------------------------------------

class TestFormattingExtraction:
    """Extractor must capture bold, italic, color, size — not just text."""

    def setup_method(self):
        self.pdf_bytes = _load_template_if_available()
        self.extractor = FormattingExtractor()
        self.doc = self.extractor.extract(self.pdf_bytes, "template")

    def test_extracts_bold_spans(self):
        """Template has ~238 bold spans. Extractor must find them."""
        bold_count = sum(
            1 for page in self.doc.pages
            for para in page.paragraphs
            for line in para.lines
            for span in line.spans
            if span.bold and span.text.strip()
        )
        # Template has 238 bold spans — we should find at least 150
        assert bold_count >= 150, f"Only {bold_count} bold spans found, expected >= 150"

    def test_extracts_italic_spans(self):
        """Template has ~91 italic spans. Extractor must find them."""
        italic_count = sum(
            1 for page in self.doc.pages
            for para in page.paragraphs
            for line in para.lines
            for span in line.spans
            if span.italic and span.text.strip()
        )
        assert italic_count >= 50, f"Only {italic_count} italic spans found, expected >= 50"

    def test_extracts_red_color_spans(self):
        """Template has 52 red (#FF0000) instruction spans."""
        red_count = sum(
            1 for page in self.doc.pages
            for para in page.paragraphs
            for line in para.lines
            for span in line.spans
            if span.color == 0xFF0000 and span.text.strip()
        )
        assert red_count >= 20, f"Only {red_count} red spans found, expected >= 20"

    def test_extracts_blue_color_spans(self):
        """Template has 47 blue link spans."""
        blue_count = sum(
            1 for page in self.doc.pages
            for para in page.paragraphs
            for line in para.lines
            for span in line.spans
            if span.color in (0x0000FF, 0x3366FF, 0x0033CC) and span.text.strip()
        )
        assert blue_count >= 15, f"Only {blue_count} blue spans found, expected >= 15"

    def test_extracts_multiple_font_sizes(self):
        """Template uses 8pt, 10pt, 11pt, 12pt, 14pt."""
        sizes = set()
        for page in self.doc.pages:
            for para in page.paragraphs:
                for line in para.lines:
                    for span in line.spans:
                        if span.text.strip():
                            sizes.add(round(span.size))
        assert len(sizes) >= 3, f"Only {len(sizes)} font sizes found: {sizes}"
        assert 11 in sizes, "Primary size 11pt not found"

    def test_extracts_indent_levels(self):
        """Template has 8+ indent levels for nested clauses."""
        indent_levels = set()
        for page in self.doc.pages:
            for para in page.paragraphs:
                if para.lines:
                    indent = round(para.lines[0].indent / 10) * 10
                    indent_levels.add(indent)
        assert len(indent_levels) >= 4, f"Only {len(indent_levels)} indent levels found"


# ---------------------------------------------------------------------------
# Test 2: DOCX Renderer must preserve ALL span formatting
# ---------------------------------------------------------------------------

class TestDOCXRendererFidelity:
    """DOCX output must preserve bold, italic, color from the IR."""

    def setup_method(self):
        self.pdf_bytes = _load_template_if_available()
        self.extractor = FormattingExtractor()
        self.renderer = DOCXRenderer()
        self.doc = self.extractor.extract(self.pdf_bytes, "template")

    def test_docx_preserves_bold(self):
        """Bold spans in IR must appear as bold runs in DOCX."""
        from docx import Document as DocxDocument
        import io

        docx_bytes = self.renderer.render(self.doc)
        word_doc = DocxDocument(io.BytesIO(docx_bytes))

        bold_runs = sum(
            1 for para in word_doc.paragraphs
            for run in para.runs
            if run.font.bold and run.text.strip()
        )
        # Template has ~238 bold spans — DOCX should have at least 100 bold runs
        assert bold_runs >= 100, f"Only {bold_runs} bold runs in DOCX, expected >= 100"

    def test_docx_preserves_italic(self):
        """Italic spans in IR must appear as italic runs in DOCX."""
        from docx import Document as DocxDocument
        import io

        docx_bytes = self.renderer.render(self.doc)
        word_doc = DocxDocument(io.BytesIO(docx_bytes))

        italic_runs = sum(
            1 for para in word_doc.paragraphs
            for run in para.runs
            if run.font.italic and run.text.strip()
        )
        assert italic_runs >= 30, f"Only {italic_runs} italic runs in DOCX, expected >= 30"

    def test_docx_preserves_font_name(self):
        """Primary font (Arial) must be set on runs."""
        from docx import Document as DocxDocument
        import io

        docx_bytes = self.renderer.render(self.doc)
        word_doc = DocxDocument(io.BytesIO(docx_bytes))

        arial_runs = sum(
            1 for para in word_doc.paragraphs
            for run in para.runs
            if run.font.name and "Arial" in run.font.name
        )
        total_runs = sum(len(para.runs) for para in word_doc.paragraphs)
        assert arial_runs > total_runs * 0.5, f"Only {arial_runs}/{total_runs} runs have Arial font"

    def test_docx_preserves_colors(self):
        """Non-black colors (red, blue) must appear in DOCX runs."""
        from docx import Document as DocxDocument
        from docx.shared import RGBColor
        import io

        docx_bytes = self.renderer.render(self.doc)
        word_doc = DocxDocument(io.BytesIO(docx_bytes))

        colored_runs = sum(
            1 for para in word_doc.paragraphs
            for run in para.runs
            if run.font.color and run.font.color.rgb and run.font.color.rgb != RGBColor(0, 0, 0)
        )
        assert colored_runs >= 10, f"Only {colored_runs} colored runs in DOCX, expected >= 10"

    def test_docx_has_correct_page_size(self):
        """Output must be A4."""
        from docx import Document as DocxDocument
        import io

        docx_bytes = self.renderer.render(self.doc)
        word_doc = DocxDocument(io.BytesIO(docx_bytes))

        section = word_doc.sections[0]
        # A4 = 8.27 x 11.69 inches (within 0.1 tolerance)
        assert abs(section.page_width.inches - 8.27) < 0.2, f"Width: {section.page_width.inches}"
        assert abs(section.page_height.inches - 11.69) < 0.2, f"Height: {section.page_height.inches}"


# ---------------------------------------------------------------------------
# Test 3: Run-on word detection must not false-positive on compounds
# ---------------------------------------------------------------------------

class TestRunonWordAccuracy:
    """Run-on detector must not flag legitimate compound words."""

    def test_hyphenated_words_not_flagged(self):
        """Hyphenated compounds are legitimate, not run-ons."""
        text = "This is a Placebo-Controlled double-blind study."
        issues = detect_runon_words(text)
        flagged_words = [w for w, _, _ in issues]
        assert "Placebo-Controlled" not in flagged_words
        assert "double-blind" not in flagged_words

    def test_slashed_words_not_flagged(self):
        text = "Demographics/Medical history was reviewed."
        issues = detect_runon_words(text)
        flagged_words = [w for w, _, _ in issues]
        assert "Demographics/Medical" not in flagged_words

    def test_scientific_identifiers_not_flagged(self):
        text = "SARS-CoV-2-specific immunity was measured."
        issues = detect_runon_words(text)
        flagged_words = [w for w, _, _ in issues]
        assert "SARS-CoV-2-specific" not in flagged_words

    def test_cannot_not_flagged(self):
        text = "The subject cannot participate."
        issues = detect_runon_words(text)
        flagged_words = [w for w, _, _ in issues]
        assert "cannot" not in flagged_words

    def test_real_runon_is_flagged(self):
        """Actual concatenated words should be detected."""
        text = "This documentContainsAVeryLongRunonWordProblem here."
        issues = detect_runon_words(text)
        assert len(issues) >= 1, "Should detect at least one run-on"


# ---------------------------------------------------------------------------
# Test 4: Formula detector accuracy
# ---------------------------------------------------------------------------

class TestFormulaDetectorAccuracy:
    """Formula detector must annotate correctly without false positives."""

    def setup_method(self):
        self.detector = FormulaDetector()

    def test_chemical_co2(self):
        formulas = self.detector.detect("CO2 levels were measured")
        assert any(f.formula_type == "chemical" and "CO" in f.original_text for f in formulas)
        assert any("<sub>2</sub>" in f.html_text for f in formulas)

    def test_pk_cmax(self):
        formulas = self.detector.detect("Cmax was 45 ng/mL")
        assert any(f.formula_type == "pk" for f in formulas)
        assert any("<sub>max</sub>" in f.html_text for f in formulas)

    def test_dosing_mg_m2(self):
        formulas = self.detector.detect("Dose was 200 mg/m2")
        assert any(f.formula_type == "dosing" for f in formulas)
        assert any("<sup>2</sup>" in f.html_text for f in formulas)

    def test_no_false_positive_on_normal_text(self):
        formulas = self.detector.detect("The patient was enrolled in the study on Day 1.")
        assert len(formulas) == 0, f"False positives: {[f.original_text for f in formulas]}"


# ---------------------------------------------------------------------------
# Test 5: Fidelity checker score sanity
# ---------------------------------------------------------------------------

class TestFidelityCheckerScoring:
    """Fidelity checker must produce reasonable scores."""

    def test_clean_document_scores_high(self):
        """A well-formatted document should score above 50."""
        pdf_bytes = _load_template_if_available()
        checker = DocumentFidelityChecker()
        report = checker.check(pdf_bytes, "template")
        assert report.score >= 50, f"Clean template scored only {report.score}"

    def test_no_false_positive_runons_on_template(self):
        """Template should have < 5 run-on word detections."""
        pdf_bytes = _load_template_if_available()
        checker = DocumentFidelityChecker()
        report = checker.check(pdf_bytes, "template")
        runon_count = sum(1 for i in report.issues if i.category == "runon_word")
        assert runon_count < 5, f"Template has {runon_count} false positive run-ons"
