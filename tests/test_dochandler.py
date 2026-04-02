"""
Comprehensive tests for the Document Handling Module.

Tests cover: ingestors, renderers, round-trip fidelity, edge cases,
formula detection, and cross-format conversion.

Target: >40 tests, all passing.
"""

import json
import pytest
from pathlib import Path

from src.formatter import DocHandler, FormattedDocument
from src.formatter.extractor import (
    FormattingExtractor,
    FormattedDocument,
    FormattedPage,
    FormattedParagraph,
    FormattedLine,
    FormattedSpan,
    FormattedTable,
    FormattedTableCell,
)
from src.formatter.ingest.html_ingestor import HTMLIngestor
from src.formatter.ingest.markdown_ingestor import MarkdownIngestor
from src.formatter.ingest.text_ingestor import TextIngestor
from src.formatter.render.html_renderer import HTMLRenderer
from src.formatter.render.markdown_renderer import MarkdownRenderer
from src.formatter.render.text_renderer import TextRenderer
from src.formatter.render.json_renderer import JSONRenderer
from src.formatter.docx_renderer import DOCXRenderer
from src.formatter.formula_detector import FormulaDetector
from src.formatter.fidelity_checker import DocumentFidelityChecker, detect_runon_words


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_doc(paragraphs=None, tables=None):
    """Create a simple FormattedDocument for testing."""
    paras = paragraphs or [
        FormattedParagraph(
            style="heading1",
            lines=[FormattedLine(spans=[FormattedSpan(
                text="Test Heading", x0=0, y0=0, x1=100, y1=14,
                bold=True, size=14.0, font="Arial",
            )])]
        ),
        FormattedParagraph(
            style="body",
            lines=[FormattedLine(spans=[
                FormattedSpan(text="Normal ", x0=0, y0=20, x1=50, y1=32, size=11.0, font="Arial"),
                FormattedSpan(text="bold", x0=50, y0=20, x1=80, y1=32, bold=True, size=11.0, font="Arial"),
                FormattedSpan(text=" and ", x0=80, y0=20, x1=100, y1=32, size=11.0, font="Arial"),
                FormattedSpan(text="italic", x0=100, y0=20, x1=130, y1=32, italic=True, size=11.0, font="Arial"),
                FormattedSpan(text=" text.", x0=130, y0=20, x1=160, y1=32, size=11.0, font="Arial"),
            ])]
        ),
    ]
    page = FormattedPage(
        page_number=0, width=595.3, height=842.0,
        paragraphs=paras,
        tables=tables or [],
    )
    return FormattedDocument(filename="test.pdf", pages=[page])


def _make_table_doc():
    """Create a doc with a table."""
    table = FormattedTable(
        num_rows=3, num_cols=2,
        rows=[
            [FormattedTableCell(text="Name", row=0, col=0, is_header=True, bold=True),
             FormattedTableCell(text="Value", row=0, col=1, is_header=True, bold=True)],
            [FormattedTableCell(text="Protocol", row=1, col=0),
             FormattedTableCell(text="C4591001", row=1, col=1)],
            [FormattedTableCell(text="Phase", row=2, col=0),
             FormattedTableCell(text="Phase 3", row=2, col=1)],
        ]
    )
    return _make_doc(tables=[table])


# ===========================================================================
# HTML Ingestor Tests
# ===========================================================================

class TestHTMLIngestor:
    def setup_method(self):
        self.ingestor = HTMLIngestor()

    def test_basic_paragraph(self):
        doc = self.ingestor.ingest("<p>Hello world</p>")
        assert doc.total_paragraphs >= 1
        text = doc.pages[0].paragraphs[0].text
        assert "Hello world" in text

    def test_heading_levels(self):
        html = "<h1>H1</h1><h2>H2</h2><h3>H3</h3>"
        doc = self.ingestor.ingest(html)
        styles = [p.style for p in doc.pages[0].paragraphs]
        assert "heading1" in styles
        assert "heading2" in styles
        assert "heading3" in styles

    def test_bold_italic(self):
        html = "<p><strong>bold</strong> and <em>italic</em></p>"
        doc = self.ingestor.ingest(html)
        spans = [s for p in doc.pages[0].paragraphs for l in p.lines for s in l.spans]
        bold_spans = [s for s in spans if s.bold]
        italic_spans = [s for s in spans if s.italic]
        assert len(bold_spans) >= 1
        assert len(italic_spans) >= 1

    def test_table_extraction(self):
        html = "<table><tr><th>A</th><th>B</th></tr><tr><td>1</td><td>2</td></tr></table>"
        doc = self.ingestor.ingest(html)
        assert len(doc.pages[0].tables) >= 1
        table = doc.pages[0].tables[0]
        assert table.num_rows == 2
        assert table.num_cols == 2

    def test_list_items(self):
        html = "<ul><li>Item 1</li><li>Item 2</li></ul>"
        doc = self.ingestor.ingest(html)
        list_paras = [p for p in doc.pages[0].paragraphs if "list" in p.style]
        assert len(list_paras) >= 2

    def test_superscript_subscript(self):
        html = "<p>CO<sub>2</sub> and 10<sup>6</sup></p>"
        doc = self.ingestor.ingest(html)
        spans = [s for p in doc.pages[0].paragraphs for l in p.lines for s in l.spans]
        sub_spans = [s for s in spans if s.subscript]
        sup_spans = [s for s in spans if s.superscript]
        assert len(sub_spans) >= 1
        assert len(sup_spans) >= 1

    def test_inline_color(self):
        html = '<p><span style="color:red">red text</span></p>'
        doc = self.ingestor.ingest(html)
        spans = [s for p in doc.pages[0].paragraphs for l in p.lines for s in l.spans]
        colored = [s for s in spans if s.color != 0]
        assert len(colored) >= 1

    def test_empty_input(self):
        doc = self.ingestor.ingest("")
        assert doc.total_paragraphs == 0

    def test_nested_formatting(self):
        html = "<p><strong><em>bold italic</em></strong></p>"
        doc = self.ingestor.ingest(html)
        spans = [s for p in doc.pages[0].paragraphs for l in p.lines for s in l.spans]
        both = [s for s in spans if s.bold and s.italic]
        assert len(both) >= 1


# ===========================================================================
# Markdown Ingestor Tests
# ===========================================================================

class TestMarkdownIngestor:
    def setup_method(self):
        self.ingestor = MarkdownIngestor()

    def test_headings(self):
        doc = self.ingestor.ingest("# H1\n\n## H2\n\n### H3")
        styles = [p.style for p in doc.pages[0].paragraphs]
        assert "heading1" in styles
        assert "heading2" in styles

    def test_bold_italic(self):
        doc = self.ingestor.ingest("**bold** and *italic* text")
        spans = [s for p in doc.pages[0].paragraphs for l in p.lines for s in l.spans]
        assert any(s.bold for s in spans)
        assert any(s.italic for s in spans)

    def test_bullet_list(self):
        doc = self.ingestor.ingest("- Item A\n- Item B\n- Item C")
        list_paras = [p for p in doc.pages[0].paragraphs if "list" in p.style]
        assert len(list_paras) >= 2

    def test_numbered_list(self):
        doc = self.ingestor.ingest("1. First\n2. Second\n3. Third")
        list_paras = [p for p in doc.pages[0].paragraphs if "list" in p.style]
        assert len(list_paras) >= 2

    def test_table(self):
        md = "| Col A | Col B |\n|-------|-------|\n| 1 | 2 |\n| 3 | 4 |"
        doc = self.ingestor.ingest(md)
        assert len(doc.pages[0].tables) >= 1

    def test_empty_input(self):
        doc = self.ingestor.ingest("")
        assert doc.total_paragraphs == 0

    def test_code_span(self):
        doc = self.ingestor.ingest("Use `print()` function")
        # Code spans should be preserved as text
        text = doc.pages[0].paragraphs[0].text
        assert "print()" in text


# ===========================================================================
# Text Ingestor Tests
# ===========================================================================

class TestTextIngestor:
    def setup_method(self):
        self.ingestor = TextIngestor()

    def test_paragraph_splitting(self):
        doc = self.ingestor.ingest("Para 1\n\nPara 2\n\nPara 3")
        assert doc.total_paragraphs >= 3

    def test_heading_detection(self):
        doc = self.ingestor.ingest("1. Introduction\n\nSome body text here.")
        styles = [p.style for p in doc.pages[0].paragraphs]
        assert any("heading" in s for s in styles)

    def test_bullet_detection(self):
        doc = self.ingestor.ingest("- Item 1\n- Item 2")
        styles = [p.style for p in doc.pages[0].paragraphs]
        assert any("list" in s for s in styles)

    def test_empty_input(self):
        doc = self.ingestor.ingest("")
        assert doc.total_paragraphs == 0


# ===========================================================================
# HTML Renderer Tests
# ===========================================================================

class TestHTMLRenderer:
    def setup_method(self):
        self.renderer = HTMLRenderer()

    def test_bold_in_output(self):
        doc = _make_doc()
        html = self.renderer.render(doc)
        assert "<strong>" in html

    def test_italic_in_output(self):
        doc = _make_doc()
        html = self.renderer.render(doc)
        assert "<em>" in html

    def test_heading_tag(self):
        doc = _make_doc()
        html = self.renderer.render(doc)
        assert "<h1>" in html or "<h1 " in html

    def test_table_rendering(self):
        doc = _make_table_doc()
        html = self.renderer.render(doc)
        assert "<table" in html
        assert "<th" in html or "<td" in html


# ===========================================================================
# Markdown Renderer Tests
# ===========================================================================

class TestMarkdownRenderer:
    def setup_method(self):
        self.renderer = MarkdownRenderer()

    def test_heading_markers(self):
        doc = _make_doc()
        md = self.renderer.render(doc)
        assert md.startswith("#") or "# " in md

    def test_bold_markers(self):
        doc = _make_doc()
        md = self.renderer.render(doc)
        assert "**" in md

    def test_italic_markers(self):
        doc = _make_doc()
        md = self.renderer.render(doc)
        assert "*italic*" in md or "_italic_" in md

    def test_table_rendering(self):
        doc = _make_table_doc()
        md = self.renderer.render(doc)
        assert "|" in md


# ===========================================================================
# Text Renderer Tests
# ===========================================================================

class TestTextRenderer:
    def setup_method(self):
        self.renderer = TextRenderer()

    def test_strips_formatting(self):
        doc = _make_doc()
        text = self.renderer.render(doc)
        assert "<strong>" not in text
        assert "**" not in text
        assert "bold" in text

    def test_preserves_content(self):
        doc = _make_doc()
        text = self.renderer.render(doc)
        assert "Test Heading" in text or "TEST HEADING" in text
        assert "italic" in text


# ===========================================================================
# JSON Renderer Tests
# ===========================================================================

class TestJSONRenderer:
    def setup_method(self):
        self.renderer = JSONRenderer()

    def test_valid_json(self):
        doc = _make_doc()
        result = self.renderer.render(doc)
        data = json.loads(result)
        assert "pages" in data
        assert "metadata" in data

    def test_preserves_span_attributes(self):
        doc = _make_doc()
        result = self.renderer.render(doc)
        data = json.loads(result)
        spans = data["pages"][0]["paragraphs"][1]["lines"][0]["spans"]
        bold_span = [s for s in spans if s.get("bold")]
        assert len(bold_span) >= 1

    def test_includes_tables(self):
        doc = _make_table_doc()
        result = self.renderer.render(doc)
        data = json.loads(result)
        assert len(data["pages"][0].get("tables", [])) >= 1


# ===========================================================================
# DOCX Renderer Tests
# ===========================================================================

class TestDOCXRenderer:
    def setup_method(self):
        self.renderer = DOCXRenderer()

    def test_produces_valid_docx(self):
        doc = _make_doc()
        result = self.renderer.render(doc)
        assert isinstance(result, bytes)
        # DOCX files start with PK (ZIP header)
        assert result[:2] == b"PK"

    def test_preserves_bold(self):
        from docx import Document
        import io
        doc = _make_doc()
        result = self.renderer.render(doc)
        word_doc = Document(io.BytesIO(result))
        bold_runs = [r for p in word_doc.paragraphs for r in p.runs if r.font.bold]
        assert len(bold_runs) >= 1

    def test_preserves_italic(self):
        from docx import Document
        import io
        doc = _make_doc()
        result = self.renderer.render(doc)
        word_doc = Document(io.BytesIO(result))
        italic_runs = [r for p in word_doc.paragraphs for r in p.runs if r.font.italic]
        assert len(italic_runs) >= 1

    def test_renders_tables(self):
        from docx import Document
        import io
        doc = _make_table_doc()
        result = self.renderer.render(doc)
        word_doc = Document(io.BytesIO(result))
        assert len(word_doc.tables) >= 1


# ===========================================================================
# DocHandler Integration Tests
# ===========================================================================

class TestDocHandler:
    def setup_method(self):
        self.handler = DocHandler()

    def test_html_to_markdown(self):
        html = "<h1>Title</h1><p><strong>Bold</strong> paragraph</p>"
        md = self.handler.convert(html, "html", "markdown")
        assert "#" in md
        assert "**" in md

    def test_markdown_to_html(self):
        md = "# Title\n\n**Bold** paragraph"
        html = self.handler.convert(md, "markdown", "html")
        assert "<h1" in html or "<strong>" in html

    def test_html_to_text(self):
        html = "<h1>Title</h1><p>Body text</p>"
        text = self.handler.convert(html, "html", "text")
        assert "Title" in text or "TITLE" in text
        assert "Body text" in text

    def test_html_to_json(self):
        html = "<p>Test paragraph</p>"
        result = self.handler.convert(html, "html", "json")
        data = json.loads(result)
        assert "pages" in data

    def test_html_to_docx(self):
        html = "<h1>Title</h1><p><strong>Bold</strong></p>"
        result = self.handler.convert(html, "html", "docx")
        assert isinstance(result, bytes)
        assert result[:2] == b"PK"

    def test_text_to_html(self):
        text = "1. Introduction\n\nSome body text."
        html = self.handler.convert(text, "text", "html")
        assert "<" in html  # has HTML tags

    def test_unsupported_input_format(self):
        with pytest.raises(ValueError, match="Unsupported input"):
            self.handler.ingest("content", "xlsx")

    def test_unsupported_output_format(self):
        doc = _make_doc()
        with pytest.raises(ValueError, match="Unsupported output"):
            self.handler.render(doc, "pdf")

    def test_round_trip_html(self):
        """HTML -> IR -> HTML should preserve content."""
        original = "<h1>Title</h1><p><strong>Bold</strong> and <em>italic</em> text</p>"
        doc = self.handler.ingest(original, "html")
        output = self.handler.render(doc, "html")
        assert "Bold" in output
        assert "italic" in output
        assert "<strong>" in output or "<b>" in output

    def test_round_trip_markdown(self):
        """Markdown -> IR -> Markdown should preserve content."""
        original = "# Title\n\n**Bold** and *italic* text"
        doc = self.handler.ingest(original, "markdown")
        output = self.handler.render(doc, "markdown")
        assert "Bold" in output
        assert "italic" in output


# ===========================================================================
# Formula Detector Tests (edge cases)
# ===========================================================================

class TestFormulaEdgeCases:
    def setup_method(self):
        self.detector = FormulaDetector()

    def test_multiple_formulas_in_one_sentence(self):
        text = "CO2 levels, HbA1c, and AUC0-inf were measured at 200 mg/m2."
        formulas = self.detector.detect(text)
        types = {f.formula_type for f in formulas}
        assert "chemical" in types
        assert "pk" in types
        assert "dosing" in types

    def test_preserves_surrounding_text(self):
        text = "The CO2 level was normal."
        result = self.detector.annotate_html(text)
        assert "The " in result
        assert " level was normal." in result
        assert "<sub>2</sub>" in result

    def test_statistical_patterns(self):
        text = "HR 0.67 (95% CI: 0.45-0.99), p < 0.001"
        formulas = self.detector.detect(text)
        stat_formulas = [f for f in formulas if f.formula_type == "statistical"]
        assert len(stat_formulas) >= 2

    def test_pk_tmax_tlag(self):
        formulas = self.detector.detect("tmax was 2h and t1/2 was 6h")
        pk = [f for f in formulas if f.formula_type == "pk"]
        assert len(pk) >= 2


# ===========================================================================
# Run-on Word Edge Cases
# ===========================================================================

class TestRunonEdgeCases:
    def test_url_not_flagged(self):
        issues = detect_runon_words("Visit https://example.com/page for details")
        assert len(issues) == 0

    def test_email_not_flagged(self):
        issues = detect_runon_words("Contact user@example.com for help")
        assert len(issues) == 0

    def test_abbreviation_not_flagged(self):
        issues = detect_runon_words("The U.S.A. and U.K. participated")
        assert len(issues) == 0

    def test_version_number_not_flagged(self):
        issues = detect_runon_words("Protocol v2.3.1 was approved")
        assert len(issues) == 0
