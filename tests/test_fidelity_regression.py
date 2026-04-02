"""
Fidelity Regression Suite — span-level metrics as CI quality gates.

Every PR must pass these tests. No metric can drop below baseline.

Tested against the CTSA template (31-page legal document) and
the P-14 clinical protocol (193-page vaccine trial).
"""

import io
import pytest
from pathlib import Path

from src.formatter.analyze.span_forensics import SpanForensics
from src.formatter.extractor import (
    FormattingExtractor,
    FormattedDocument,
    FormattedPage,
    FormattedParagraph,
    FormattedLine,
    FormattedSpan,
    FormattedTable,
    FormattedTableCell,
)
from src.formatter.docx_renderer import DOCXRenderer


CTSA_PATH = Path("C:/Users/kapil/Documents/051 ANIMATE CTSA template_15-05-18_22062018_0.pdf")
P14_PATH = Path("data_seed/pdfs/P-14.pdf")


def _load_if_available(path: Path) -> bytes:
    if not path.exists():
        pytest.skip(f"Test document not available: {path}")
    return path.read_bytes()


# ===========================================================================
# Span-Level Fidelity Gates (no metric may drop below baseline)
# ===========================================================================

class TestCTSAFidelity:
    """Fidelity gates for the CTSA legal template."""

    def setup_method(self):
        self.pdf = _load_if_available(CTSA_PATH)
        self.forensics = SpanForensics()
        self.report = self.forensics.analyze_pdf(self.pdf, "CTSA")

    def test_bold_fidelity_above_93(self):
        assert self.report.fidelity("bold") >= 0.93

    def test_italic_fidelity_above_90(self):
        assert self.report.fidelity("italic") >= 0.90

    def test_colored_fidelity_above_98(self):
        assert self.report.fidelity("colored") >= 0.98

    def test_end_to_end_bold_above_93(self):
        assert self.report.render_fidelity("bold") >= 0.93

    def test_end_to_end_italic_above_90(self):
        assert self.report.render_fidelity("italic") >= 0.90

    def test_end_to_end_colored_above_98(self):
        assert self.report.render_fidelity("colored") >= 0.98


class TestP14Fidelity:
    """Fidelity gates for the P-14 clinical protocol."""

    def setup_method(self):
        self.pdf = _load_if_available(P14_PATH)
        self.forensics = SpanForensics()
        self.report = self.forensics.analyze_pdf(self.pdf, "P-14")

    def test_bold_fidelity_above_93(self):
        assert self.report.fidelity("bold") >= 0.93

    def test_italic_fidelity_above_90(self):
        assert self.report.fidelity("italic") >= 0.90

    def test_underline_fidelity_above_85(self):
        assert self.report.fidelity("underline") >= 0.85

    def test_superscript_fidelity_above_90(self):
        assert self.report.fidelity("superscript") >= 0.90

    def test_subscript_fidelity_above_90(self):
        assert self.report.fidelity("subscript") >= 0.90

    def test_colored_fidelity_above_98(self):
        assert self.report.fidelity("colored") >= 0.98

    def test_end_to_end_bold(self):
        assert self.report.render_fidelity("bold") >= 0.93

    def test_end_to_end_underline(self):
        assert self.report.render_fidelity("underline") >= 0.85


# ===========================================================================
# Round-Trip Attribute Preservation (synthetic document)
# ===========================================================================

class TestSyntheticRoundTrip:
    """Every IR attribute must survive render → re-ingest."""

    def setup_method(self):
        """Create a synthetic IR document with every attribute set."""
        self.doc = FormattedDocument(
            filename="synthetic_test",
            pages=[FormattedPage(
                page_number=0, width=595.3, height=842.0,
                paragraphs=[
                    # Heading
                    FormattedParagraph(
                        style="heading1",
                        lines=[FormattedLine(spans=[FormattedSpan(
                            text="Test Heading", x0=72, y0=72, x1=300, y1=86,
                            font="Arial", size=16.0, bold=True,
                        )])]
                    ),
                    # Body with bold
                    FormattedParagraph(
                        style="body",
                        lines=[FormattedLine(spans=[
                            FormattedSpan(text="Normal ", x0=72, y0=100, x1=120, y1=112, font="Arial", size=11.0),
                            FormattedSpan(text="bold", x0=120, y0=100, x1=150, y1=112, font="Arial", size=11.0, bold=True),
                            FormattedSpan(text=" and ", x0=150, y0=100, x1=170, y1=112, font="Arial", size=11.0),
                            FormattedSpan(text="italic", x0=170, y0=100, x1=200, y1=112, font="Arial", size=11.0, italic=True),
                            FormattedSpan(text=" text.", x0=200, y0=100, x1=240, y1=112, font="Arial", size=11.0),
                        ])]
                    ),
                    # Colored text
                    FormattedParagraph(
                        style="body",
                        lines=[FormattedLine(spans=[
                            FormattedSpan(text="Red text", x0=72, y0=130, x1=150, y1=142, font="Arial", size=11.0, color=0xFF0000),
                            FormattedSpan(text=" blue text", x0=150, y0=130, x1=230, y1=142, font="Arial", size=11.0, color=0x0000FF),
                        ])]
                    ),
                    # Super/subscript
                    FormattedParagraph(
                        style="body",
                        lines=[FormattedLine(spans=[
                            FormattedSpan(text="CO", x0=72, y0=160, x1=90, y1=172, font="Arial", size=11.0),
                            FormattedSpan(text="2", x0=90, y0=160, x1=96, y1=172, font="Arial", size=7.0, subscript=True),
                            FormattedSpan(text=" and 10", x0=96, y0=160, x1=140, y1=172, font="Arial", size=11.0),
                            FormattedSpan(text="6", x0=140, y0=160, x1=146, y1=172, font="Arial", size=7.0, superscript=True),
                        ])]
                    ),
                ],
                tables=[FormattedTable(
                    num_rows=2, num_cols=2,
                    rows=[
                        [FormattedTableCell(text="Header A", row=0, col=0, is_header=True, bold=True),
                         FormattedTableCell(text="Header B", row=0, col=1, is_header=True, bold=True)],
                        [FormattedTableCell(text="Cell 1", row=1, col=0),
                         FormattedTableCell(text="Cell 2", row=1, col=1)],
                    ]
                )],
            )]
        )

    def test_docx_round_trip_bold(self):
        """Bold spans must survive DOCX round-trip."""
        from docx import Document
        renderer = DOCXRenderer()
        docx_bytes = renderer.render(self.doc)
        word = Document(io.BytesIO(docx_bytes))
        bold_runs = [r for p in word.paragraphs for r in p.runs if r.font.bold and r.text.strip()]
        assert len(bold_runs) >= 2  # heading + inline bold

    def test_docx_round_trip_italic(self):
        from docx import Document
        renderer = DOCXRenderer()
        docx_bytes = renderer.render(self.doc)
        word = Document(io.BytesIO(docx_bytes))
        italic_runs = [r for p in word.paragraphs for r in p.runs if r.font.italic and r.text.strip()]
        assert len(italic_runs) >= 1

    def test_docx_round_trip_color(self):
        from docx import Document
        from docx.shared import RGBColor
        renderer = DOCXRenderer()
        docx_bytes = renderer.render(self.doc)
        word = Document(io.BytesIO(docx_bytes))
        colored = [r for p in word.paragraphs for r in p.runs
                   if r.font.color and r.font.color.rgb and r.font.color.rgb != RGBColor(0, 0, 0)]
        assert len(colored) >= 2  # red + blue

    def test_docx_round_trip_superscript(self):
        from docx import Document
        renderer = DOCXRenderer()
        docx_bytes = renderer.render(self.doc)
        word = Document(io.BytesIO(docx_bytes))
        sup = [r for p in word.paragraphs for r in p.runs if r.font.superscript]
        assert len(sup) >= 1

    def test_docx_round_trip_subscript(self):
        from docx import Document
        renderer = DOCXRenderer()
        docx_bytes = renderer.render(self.doc)
        word = Document(io.BytesIO(docx_bytes))
        sub = [r for p in word.paragraphs for r in p.runs if r.font.subscript]
        assert len(sub) >= 1

    def test_docx_round_trip_tables(self):
        from docx import Document
        renderer = DOCXRenderer()
        docx_bytes = renderer.render(self.doc)
        word = Document(io.BytesIO(docx_bytes))
        assert len(word.tables) >= 1

    def test_docx_round_trip_font_name(self):
        from docx import Document
        renderer = DOCXRenderer()
        docx_bytes = renderer.render(self.doc)
        word = Document(io.BytesIO(docx_bytes))
        arial_runs = [r for p in word.paragraphs for r in p.runs if r.font.name and "Arial" in r.font.name]
        assert len(arial_runs) >= 3
