"""
Protocol Section Parser — deterministic extraction of document structure.

Extracts the full section hierarchy from a protocol PDF:
- Section numbers (1, 1.1, 1.1.1, etc.)
- Section titles
- Page numbers
- Page ranges (start → end)

This is NOT LLM-based — it uses PyMuPDF's text extraction + regex
to deterministically parse the table of contents and section headers.
Results are 100% repeatable (same PDF → same output every time).

Usage:
    parser = SectionParser()
    sections = parser.parse(pdf_bytes)
    # sections[0] = Section(number="1", title="INTRODUCTION", page=8, ...)
    # parser.find("4.1") → Section(number="4.1", title="STUDY DESIGN", page=23)
"""

from __future__ import annotations

import re
import logging
from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF

try:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)


def _looks_like_equation(text: str) -> bool:
    """Heuristic: does this line look like a mathematical equation?

    Catches:
    - Sample size formulas: n = (Z_α/2 + Z_β)² × 2σ² / δ²
    - Confidence intervals: CI = p̂ ± Z × √(p̂(1-p̂)/n)
    - Hazard ratios: HR = exp(β)
    - Summation/product notation
    - Greek letters mixed with operators
    """
    if len(text) < 3 or len(text) > 200:
        return False

    # Count math-like characters
    math_chars = set("=±×÷√∑∏∫≤≥≠≈∈∞αβγδεζηθλμσφχψω²³⁻¹")
    math_count = sum(1 for c in text if c in math_chars)

    # Standard equation patterns
    equation_patterns = [
        r"[=≈≤≥]",       # Has equals/comparison
        r"[±√∑∏∫]",      # Has math operators
        r"[αβγδσμ]",     # Has Greek letters
        r"\^[0-9\{]",    # Has superscripts
        r"_[0-9\{]",     # Has subscripts
        r"\\frac",        # LaTeX fraction
        r"\bexp\b",       # Exponential
        r"\blog\b",       # Logarithm
        r"\bln\b",        # Natural log
        r"[²³⁻¹⁺]",     # Unicode superscripts
        r"\bCI\s*=",      # Confidence interval formula
        r"\bn\s*=",       # Sample size formula
        r"\bHR\s*=",      # Hazard ratio
        r"\bp\s*[<>=]",   # P-value
        r"Z[_α]",        # Z-score
    ]

    import re as _re
    pattern_matches = sum(1 for p in equation_patterns if _re.search(p, text))

    # Require stronger signal: 3+ math characters OR 3+ pattern matches.
    # Previous threshold (2) triggered on common clinical text like
    # "≥12 years (Phase 2/3)" or "n = 636 patients".
    return math_count >= 3 or pattern_matches >= 3

# Matches section headers:
# Numbered: "1.", "1.1", "1.1.1", "4.2.3.1"
# Lettered: "A.", "B.1", "C.2.3"
# Roman: "I.", "II.", "III.1"
_SECTION_RE = re.compile(
    r"^(\d{1,2}(?:\.\d{1,3}){0,4})\.?\s+([A-Z][A-Za-z\s,\-/&\(\):;'\"]+)"
)
_SECTION_LETTER_RE = re.compile(
    r"^([A-Z](?:\.\d{1,3}){0,3})\.?\s+([A-Z][A-Za-z\s,\-/&\(\):;'\"]+)"
)
_SECTION_ROMAN_RE = re.compile(
    r"^((?:IX|IV|V?I{1,3}|VI{1,3}|X{1,3})(?:\.\d{1,3}){0,3})\.?\s+([A-Z][A-Za-z\s,\-/&\(\):;'\"]+)"
)

# Matches TOC entries like "1.  INTRODUCTION .................. 8"
_TOC_RE = re.compile(
    r"(\d{1,2}(?:\.\d{1,3}){0,4})\.?\s+"
    r"([A-Za-z][A-Za-z\s,\-/&\(\)]+?)"
    r"\s*[\.·\-_]{3,}\s*(\d{1,4})\s*$"
)

# Also match TOC without dots: "1.  INTRODUCTION  8"
_TOC_SIMPLE_RE = re.compile(
    r"(\d{1,2}(?:\.\d{1,3}){0,4})\.?\s+"
    r"([A-Z][A-Z\s,\-/&\(\)]{3,}?)"
    r"\s+(\d{1,4})\s*$"
)


@dataclass
class Section:
    number: str          # "1", "1.1", "4.2.3"
    title: str           # "INTRODUCTION", "Study Design"
    page: int            # 0-indexed page number
    page_display: str    # Display page number from TOC (may differ from 0-indexed)
    level: int           # Nesting depth: 1, 2, 3, 4
    end_page: int | None = None  # Last page of this section
    start_y: float = 0.0  # Y-coordinate of heading on start page (for precise clipping)
    children: list[Section] = field(default_factory=list)

    @property
    def full_title(self) -> str:
        return f"{self.number} {self.title}"

    def to_dict(self) -> dict:
        return {
            "number": self.number,
            "title": self.title,
            "page": self.page,
            "page_display": self.page_display,
            "level": self.level,
            "end_page": self.end_page,
            "start_y": self.start_y,
            "children": [c.to_dict() for c in self.children],
        }


class SectionParser:
    """Deterministic protocol section parser using PyMuPDF text extraction."""

    def parse(self, file_bytes: bytes, filename: str = "") -> list[Section]:
        """Parse all sections from a protocol PDF or DOCX.

        Auto-detects format from filename or magic bytes.
        """
        is_docx = (
            filename.lower().endswith(".docx")
            or file_bytes[:4] == b"PK\x03\x04"  # ZIP/DOCX magic bytes
        )

        if is_docx:
            return self.parse_docx(file_bytes)
        return self.parse_pdf(file_bytes)

    def parse_docx(self, docx_bytes: bytes) -> list[Section]:
        """Parse sections from a DOCX file using python-docx.

        DOCX is MUCH easier than PDF — paragraphs have style metadata
        (Heading 1, Heading 2, etc.) that directly maps to section hierarchy.
        """
        if not HAS_DOCX:
            logger.error("python-docx not installed — cannot parse DOCX")
            return []

        import io
        doc = DocxDocument(io.BytesIO(docx_bytes))
        sections: list[Section] = []
        seen = set()

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # Check if this paragraph is a heading
            style_name = (para.style.name or "").lower()
            level = 0

            if "heading 1" in style_name:
                level = 1
            elif "heading 2" in style_name:
                level = 2
            elif "heading 3" in style_name:
                level = 3
            elif "heading 4" in style_name:
                level = 4
            elif "heading 5" in style_name:
                level = 5

            if level == 0:
                # Try regex match on text for numbered sections
                match = _SECTION_RE.match(text)
                if match:
                    sec_num = match.group(1)
                    level = sec_num.count(".") + 1
                else:
                    continue

            # Extract section number from text
            match = re.match(r"^(\d{1,2}(?:\.\d{1,3}){0,4})\.?\s+(.*)", text)
            if match:
                sec_num = match.group(1)
                sec_title = match.group(2).strip()
            else:
                sec_num = ""
                sec_title = text

            key = f"{sec_num}_{sec_title[:20]}"
            if key in seen:
                continue
            seen.add(key)

            sections.append(Section(
                number=sec_num,
                title=sec_title,
                page=i,  # paragraph index (no page concept in DOCX)
                page_display=str(i + 1),
                level=level,
            ))

        logger.info(f"Parsed {len(sections)} sections from DOCX")
        # Page ranges don't apply for DOCX — use paragraph indices
        for j, s in enumerate(sections):
            if j + 1 < len(sections):
                s.end_page = sections[j + 1].page - 1
            else:
                s.end_page = len(doc.paragraphs) - 1

        return sections

    def get_section_text_docx(self, docx_bytes: bytes, section: Section) -> str:
        """Extract verbatim text from a DOCX section. 100% accurate."""
        if not HAS_DOCX:
            return ""

        import io
        doc = DocxDocument(io.BytesIO(docx_bytes))
        start = section.page
        end = section.end_page if section.end_page is not None else len(doc.paragraphs) - 1

        text_parts = []
        for i in range(start, min(end + 1, len(doc.paragraphs))):
            text_parts.append(doc.paragraphs[i].text)

        return "\n".join(text_parts).strip()

    def get_section_docx_xml(self, docx_bytes: bytes, section: Section) -> str:
        """Extract section content as raw DOCX XML — preserves equations.

        Returns the raw OpenXML paragraph elements including:
        - OMML equations (editable in Word/MathType)
        - Formatting (bold, italic, underline)
        - Lists and numbering
        - Table references

        This enables round-trip editing: extract → edit in Word → paste back.
        """
        if not HAS_DOCX:
            return ""

        import io
        from lxml import etree

        doc = DocxDocument(io.BytesIO(docx_bytes))
        start = section.page
        end = section.end_page if section.end_page is not None else len(doc.paragraphs) - 1

        xml_parts = []
        for i in range(start, min(end + 1, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            # Get the raw XML element which preserves OMML equations
            xml_str = etree.tostring(para._element, pretty_print=True).decode("utf-8")
            xml_parts.append(xml_str)

        return "\n".join(xml_parts)

    def get_section_with_equations(self, pdf_bytes: bytes, section: Section) -> dict:
        """Extract section content with equations detected and converted to LaTeX.

        For clinical protocols with statistical formulas (sample size calculations,
        confidence intervals, hazard ratios), this method:
        1. Extracts the full text
        2. Detects equation-like patterns
        3. Attempts to preserve the visual layout of multi-line formulas
        4. Wraps detected equations in LaTeX $...$ delimiters

        Returns:
            dict with 'text', 'equations' (list of detected formulas),
            and 'latex' (text with equations wrapped in LaTeX)
        """
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        start = section.page
        end = section.end_page if section.end_page is not None else start
        end = min(end + 1, doc.page_count - 1)

        equations = []
        text_with_layout = []

        for page_num in range(start, end + 1):
            page = doc[page_num]

            # Get text blocks with position info to preserve layout
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block.get("type") != 0:
                    continue
                for line_data in block.get("lines", []):
                    line_text = ""
                    has_math_font = False

                    for span in line_data.get("spans", []):
                        text = span.get("text", "")
                        font = span.get("font", "").lower()
                        # Detect math fonts (Symbol, Cambria Math, STIX, etc.)
                        if any(mf in font for mf in ["symbol", "math", "stix", "cambria"]):
                            has_math_font = True
                        line_text += text

                    line_stripped = line_text.strip()
                    if not line_stripped:
                        continue

                    # Detect equation patterns
                    is_equation = (
                        has_math_font
                        or _looks_like_equation(line_stripped)
                    )

                    if is_equation:
                        equations.append(line_stripped)
                        text_with_layout.append(f"$${line_stripped}$$")
                    else:
                        text_with_layout.append(line_stripped)

        doc.close()

        plain_text = self.get_section_text(pdf_bytes, section)

        return {
            "text": plain_text,
            "equations": equations,
            "latex": "\n".join(text_with_layout),
            "equation_count": len(equations),
        }

    def get_tables_in_section_docx(self, docx_bytes: bytes, section: Section) -> list[list[list[str]]]:
        """Extract tables from a DOCX section. Returns exact table data."""
        if not HAS_DOCX:
            return []

        import io
        doc = DocxDocument(io.BytesIO(docx_bytes))

        # DOCX tables are separate from paragraphs — find tables
        # that appear between this section's paragraph range
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)

        return tables

    def parse_pdf(self, pdf_bytes: bytes) -> list[Section]:
        """Parse all sections from a protocol PDF.

        Strategy (cascading fallback):
        1. PyMuPDF built-in TOC metadata (most reliable, instant)
        2. Parse TOC text from content pages (regex on dotted lines)
        3. Scan all pages for section headers (regex + font-size detection)
        4. If all strategies produce <5 sections, mark as low-confidence

        Returns sections sorted by page number.
        """
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_pages = doc.page_count

        # Try all strategies and pick the best result
        candidates: list[tuple[str, list[Section]]] = []

        # Strategy 1: Parse TOC from PyMuPDF's built-in TOC extraction
        toc = doc.get_toc()
        if toc and len(toc) > 5:
            sections = self._parse_fitz_toc(toc)
            if sections and len(sections) > 5:
                candidates.append(("fitz_toc", sections))
                logger.info(f"Strategy 1 (fitz TOC): {len(sections)} sections")

        # Strategy 2: Find the TOC pages and parse text
        sections = self._parse_toc_text(doc)
        if sections and len(sections) > 3:
            candidates.append(("toc_text", sections))
            logger.info(f"Strategy 2 (TOC text): {len(sections)} sections")

        # Strategy 3: Scan all pages for section headers (with font detection)
        sections = self._scan_headers(doc)
        if sections:
            # Sanity check: real protocols have 15-80 sections. If header_scan
            # returns 200+, it caught cover page noise (P-34 Roche format).
            # Demote to low priority instead of blocking entirely — it may be
            # the only candidate for non-ICH documents.
            if len(sections) > 200:
                logger.warning(
                    f"Strategy 3 produced {len(sections)} sections — likely "
                    f"cover page noise. Demoting to last resort."
                )
                candidates.append(("header_scan_noisy", sections))
            else:
                candidates.append(("header_scan", sections))
            logger.info(f"Strategy 3 (header scan): {len(sections)} sections")

        doc.close()

        # Pick the best candidate — TOC-based strategies strongly preferred
        if not candidates:
            logger.warning("No sections found by any strategy")
            return []

        # Strategy priority: fitz_toc > toc_text > header_scan > header_scan_noisy
        # TOC-based strategies have correct page numbers — strongly preferred
        # Noisy header_scan is dead last — only used if nothing else works
        priority_order = ["fitz_toc", "toc_text", "header_scan", "header_scan_noisy"]
        best_name, best_sections = None, []

        for strategy in priority_order:
            match = [(name, secs) for name, secs in candidates if name == strategy]
            if match and len(match[0][1]) >= 8:
                best_name, best_sections = match[0]
                break

        # If no strategy hit threshold, fall back to scoring
        if not best_name:
            strategy_bonus = {"fitz_toc": 500, "toc_text": 300, "header_scan": 0}
            best_name, best_sections = max(
                candidates,
                key=lambda c: self._score_sections(c[1], total_pages) + strategy_bonus.get(c[0], 0),
            )

        best_score = self._score_sections(best_sections, total_pages)
        logger.info(
            f"Selected strategy '{best_name}' with {len(best_sections)} sections "
            f"(score={best_score:.0f}, from {len(candidates)} candidates)"
        )

        # Strategy 4: If best result is weak (<10 valid sections), flag for LLM fallback
        if best_score < 30 or len(best_sections) < 5:
            logger.warning(
                f"Low-confidence section parse ({len(best_sections)} sections, "
                f"score={best_score:.0f}). LLM fallback recommended — call "
                f"parse_with_llm() for better results."
            )
            # Store the flag so callers know
            self._low_confidence = True
        else:
            self._low_confidence = False

        self._compute_page_ranges(best_sections, total_pages)
        return best_sections

    @property
    def needs_llm_fallback(self) -> bool:
        """True if the last parse produced low-confidence results."""
        return getattr(self, "_low_confidence", False)

    async def parse_with_llm(
        self,
        pdf_bytes: bytes,
        llm_client: Any = None,
        max_toc_pages: int = 8,
    ) -> list[Section]:
        """LLM-assisted section parsing — fallback when deterministic parsing fails.

        Sends the first few pages (typically TOC) to the VLM and asks it
        to extract the complete section outline. The LLM reads visual formatting
        (bold, indentation, font size) that PyMuPDF text extraction misses.

        Cost: ~$0.10 per protocol (one vision call on ~5-8 pages).
        """
        if llm_client is None:
            logger.warning("No LLM client provided for fallback parsing")
            return self.parse(pdf_bytes)

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_pages = doc.page_count

        # Render first N pages as images for the VLM
        scale = 150 / 72.0
        matrix = fitz.Matrix(scale, scale)
        images = []
        for page_num in range(min(max_toc_pages, total_pages)):
            page = doc[page_num]
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            images.append(pix.tobytes("png"))

        doc.close()

        prompt = """These are the first pages of a clinical trial protocol document.
They contain the cover page, table of contents, and/or beginning of the protocol.

Extract the COMPLETE document outline — every section and subsection with its
section number and title. Return a JSON array:

[
  {"number": "1", "title": "PROTOCOL SUMMARY", "page": 8},
  {"number": "1.1", "title": "Synopsis", "page": 8},
  {"number": "2", "title": "INTRODUCTION", "page": 15},
  ...
]

Rules:
- Include ALL sections visible in the table of contents or page headers
- Include section numbers exactly as written (1, 1.1, 1.1.1, etc.)
- Include the page number if visible in the TOC
- If no page number is visible, use 0
- Capture unnumbered sections too (appendices, lists of tables, etc.)
- Look for bold text, larger font sizes, and indentation as heading indicators

Return ONLY the JSON array."""

        try:
            raw = await llm_client.vision_json_query_multi(
                images, prompt,
                system="You are a document structure analyst. Extract ALL sections from any document type — protocols, contracts, agreements, reports. Return valid JSON only.",
                max_tokens=4096,
            )

            if not isinstance(raw, list):
                logger.warning("LLM fallback did not return a list")
                return self.parse(pdf_bytes)

            sections = []
            for item in raw:
                if not isinstance(item, dict):
                    continue
                number = str(item.get("number") or "")
                title = str(item.get("title") or "")
                page = int(item.get("page") or 0)
                if not title:
                    continue
                level = number.count(".") + 1 if number else 1
                sections.append(Section(
                    number=number,
                    title=title,
                    page=max(0, page - 1),  # Convert to 0-indexed
                    page_display=str(page) if page else "",
                    level=level,
                ))

            if sections:
                logger.info(f"LLM fallback: {len(sections)} sections extracted")
                self._compute_page_ranges(sections, total_pages)
                self._low_confidence = False
                return sections

        except Exception as e:
            logger.error(f"LLM fallback parsing failed: {e}")

        # If LLM also fails, return deterministic results
        return self.parse(pdf_bytes)

    @staticmethod
    def _score_sections(sections: list[Section], total_pages: int) -> float:
        """Score a section list for quality. Higher = better."""
        if not sections:
            return 0

        score = len(sections)  # More sections generally better

        # Bonus for having standard top-level sections
        numbers = {s.number for s in sections}
        for expected in ["1", "2", "3", "4", "5"]:
            if expected in numbers:
                score += 10

        # Bonus for having titled sections (not just numbers)
        titled = sum(1 for s in sections if len(s.title) > 5)
        score += titled * 0.5

        # Penalty for suspicious patterns (dates parsed as sections)
        months = {"january", "february", "march", "april", "may", "june",
                  "july", "august", "september", "october", "november", "december",
                  "jan", "feb", "mar", "apr", "jun", "jul", "aug", "sep", "oct", "nov", "dec"}
        date_sections = sum(
            1 for s in sections
            if s.title and s.title.split() and s.title.lower().split()[0] in months
        )
        score -= date_sections * 20  # Heavy penalty

        # Penalty for all sections on same page (likely cover page noise)
        if len(sections) > 3:
            pages_used = len({s.page for s in sections})
            if pages_used == 1:
                score -= 50  # All on one page = probably wrong

        return score

    def find(self, sections: list[Section], section_number: str) -> Section | None:
        """Find a section by its number (e.g., "4.1", "8.2.3")."""
        target = section_number.strip().rstrip(".")
        for s in self._flatten(sections):
            if s.number == target:
                return s
        return None

    def find_by_title(self, sections: list[Section], title_query: str) -> list[Section]:
        """Find sections whose title contains the query string."""
        query = title_query.lower()
        return [s for s in self._flatten(sections) if query in s.title.lower()]

    @staticmethod
    def _next_section_number(number: str) -> str | None:
        """Get the next sibling section number.

        '2.2.1' → '2.2.2', '3.1' → '3.2', '5' → '6'
        """
        if not number:
            return None
        parts = number.split(".")
        try:
            parts[-1] = str(int(parts[-1]) + 1)
            return ".".join(parts)
        except ValueError:
            return None

    def _find_heading_y(
        self, doc: fitz.Document, page_0indexed: int,
        section_number: str, title: str,
    ) -> float:
        """Find the exact Y-coordinate of a section heading on a page.

        Uses bold-font detection to pinpoint where a section starts,
        enabling precise Y-coordinate clipping instead of regex splitting.
        """
        if page_0indexed < 0 or page_0indexed >= doc.page_count:
            return 0.0

        page = doc[page_0indexed]
        blocks = page.get_text("dict")["blocks"]

        heading_prefix = f"{section_number}."
        heading_prefix_alt = f"{section_number} "

        for block in blocks:
            if "lines" not in block:
                continue
            for line in block["lines"]:
                line_text = "".join(s["text"] for s in line["spans"]).strip()
                if line_text.startswith(heading_prefix) or line_text.startswith(heading_prefix_alt):
                    is_bold = any(
                        "Bold" in s.get("font", "") or s["flags"] & 16
                        for s in line["spans"] if s["text"].strip()
                    )
                    if is_bold:
                        return line["spans"][0]["origin"][1]

            # Fallback: search by title fragment
            if title:
                for line in block.get("lines", []):
                    line_text = "".join(s["text"] for s in line["spans"]).strip()
                    if title[:30] in line_text:
                        is_bold = any(
                            "Bold" in s.get("font", "") or s["flags"] & 16
                            for s in line["spans"] if s["text"].strip()
                        )
                        if is_bold:
                            return line["spans"][0]["origin"][1]

        return 0.0

    def _find_next_heading_y(
        self, doc: fitz.Document, page_0indexed: int,
        current_section_number: str, after_y: float = 0.0,
    ) -> float:
        """Find the Y-coordinate of the next section heading on a page.

        Only considers headings BELOW after_y. Returns 99999.0 if no
        next heading found on this page.
        """
        if page_0indexed < 0 or page_0indexed >= doc.page_count:
            return 99999.0

        page = doc[page_0indexed]
        blocks = page.get_text("dict")["blocks"]
        section_re = re.compile(r"^(\d{1,2}(?:\.\d{1,2}){0,5})\.?\s+")

        candidates = []
        for block in blocks:
            if "lines" not in block:
                continue
            for line in block["lines"]:
                line_text = "".join(s["text"] for s in line["spans"]).strip()
                match = section_re.match(line_text)
                if not match:
                    continue
                found_num = match.group(1)
                if found_num == current_section_number:
                    continue
                is_bold = any(
                    "Bold" in s.get("font", "") or s["flags"] & 16
                    for s in line["spans"] if s["text"].strip()
                )
                if is_bold:
                    y = line["spans"][0]["origin"][1]
                    if y > after_y + 5:
                        candidates.append(y)

        return min(candidates) if candidates else 99999.0

    def get_section_text(
        self,
        pdf_bytes: bytes,
        section: Section,
        include_subsections: bool = True,
        preserve_formatting: bool = True,
    ) -> str:
        """Extract the EXACT text of a section from the PDF.

        Uses Y-coordinate clipping for precise extraction — on the start
        page, skips all content above the section heading. On the end page,
        stops at the next section's heading. This eliminates regex-based
        splitting bugs on shared pages.

        Args:
            section: The section to extract.
            include_subsections: If True, include subsection content too.
                If False, stop at the first subsection header.
            preserve_formatting: If True, use PyMuPDF's text with layout
                preservation (better for tables and lists).
        """
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        start = section.page
        declared_end = section.end_page if section.end_page is not None else start

        # Buffer: look up to 2 pages beyond declared end for boundary
        end = min(declared_end + 2, doc.page_count - 1)

        # Detect image-based sections (SoA tables rendered as images)
        empty_pages = 0
        actual_start = start
        for p in range(start, min(start + 10, doc.page_count)):
            page_text = doc[p].get_text("text").strip()
            if len(page_text) > 50:
                actual_start = p
                break
            empty_pages += 1

        if empty_pages >= 2:
            doc.close()
            return (
                f"[This section (pages {start+1}-{start+empty_pages}) contains "
                f"image-based content (likely a Schedule of Activities table) that "
                f"cannot be extracted as text. Use the SoA extraction pipeline "
                f"(POST /api/extract) to extract table data from these pages.]"
            )

        # Find precise Y-coordinates for clipping
        start_y = section.start_y
        if start_y == 0.0:
            start_y = self._find_heading_y(doc, actual_start, section.number, section.title)

        # Determine end page and end Y based on subsection handling
        if not include_subsections and section.children:
            first_child = section.children[0]
            end = first_child.page
            end_y = self._find_heading_y(doc, end, first_child.number, first_child.title)
            if end_y == 0.0:
                end_y = 99999.0
        else:
            # Scan each page from declared_end for the next heading to
            # prevent bleeding into the following section.
            end_y = 99999.0
            for scan_page in range(declared_end, end + 1):
                after = start_y if scan_page == actual_start else 0.0
                found_y = self._find_next_heading_y(
                    doc, scan_page, section.number, after_y=after,
                )
                if found_y < 99999.0:
                    end = scan_page
                    end_y = found_y
                    break

            # Same-page sibling boundary detection
            if end_y >= 99999.0:
                next_num = self._next_section_number(section.number)
                if next_num:
                    for scan_page in range(actual_start, end + 1):
                        sibling_y = self._find_heading_y(
                            doc, scan_page, next_num, ""
                        )
                        if sibling_y > start_y + 5:
                            end = scan_page
                            end_y = sibling_y
                            break

        # Extract text with Y-coordinate clipping per page
        lines = []
        header_footer_re = [
            re.compile(r"^Page \d+", re.IGNORECASE),
            re.compile(r"^Confidential$", re.IGNORECASE),
            re.compile(r"^\d{1,4}$"),  # Standalone page numbers
        ]

        for page_num in range(actual_start, end + 1):
            page = doc[page_num]
            blocks = page.get_text("dict")["blocks"]

            page_start_y = start_y if page_num == actual_start else 0.0
            page_end_y = end_y if page_num == end else 99999.0

            for block in blocks:
                if "lines" not in block:
                    continue
                for line in block["lines"]:
                    spans = line.get("spans", [])
                    if not spans:
                        continue

                    y = spans[0]["origin"][1]

                    # Y-coordinate clipping
                    if y < page_start_y - 2:
                        continue
                    if y > page_end_y - 2:
                        continue

                    line_text = "".join(s["text"] for s in spans).strip()
                    if not line_text:
                        continue

                    # Skip headers/footers
                    skip = False
                    for pat in header_footer_re:
                        if pat.match(line_text):
                            skip = True
                            break
                    if skip:
                        continue

                    if "Continued on table" in line_text:
                        continue

                    lines.append(line_text)

        doc.close()
        return "\n".join(lines).strip()

    def get_section_html(self, pdf_bytes: bytes, section: Section) -> str:
        """Extract section content as HTML — preserves formatting for copy-paste.

        Uses Y-coordinate clipping for precise extraction.

        Returns HTML with:
        - Paragraphs as <p> tags
        - Bold text preserved
        - Tables preserved if present
        """
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        start = section.page
        end = section.end_page if section.end_page is not None else start
        end = min(end + 1, doc.page_count - 1)

        # Find precise Y-coordinates
        start_y = section.start_y
        if start_y == 0.0:
            start_y = self._find_heading_y(doc, start, section.number, section.title)
        end_y = self._find_next_heading_y(
            doc, end, section.number,
            after_y=start_y if end == start else 0.0,
        )

        html_parts = []
        for page_num in range(start, end + 1):
            page = doc[page_num]
            page_start_y = start_y if page_num == start else 0.0
            page_end_y = end_y if page_num == end else 99999.0

            try:
                blocks = page.get_text("dict")["blocks"]
                for block in blocks:
                    if block.get("type") == 0:  # Text block
                        for line_data in block.get("lines", []):
                            spans = line_data.get("spans", [])
                            if not spans:
                                continue
                            y = spans[0]["origin"][1]
                            if y < page_start_y - 2 or y > page_end_y - 2:
                                continue

                            line_html = ""
                            for span in spans:
                                text = span.get("text", "")
                                if not text.strip():
                                    continue
                                flags = span.get("flags", 0)
                                is_bold = flags & (1 << 4)
                                is_italic = flags & (1 << 1)
                                if is_bold:
                                    text = f"<strong>{text}</strong>"
                                if is_italic:
                                    text = f"<em>{text}</em>"
                                line_html += text
                            if line_html.strip():
                                html_parts.append(f"<p>{line_html}</p>")
            except Exception:
                html_parts.append(f"<p>{page.get_text('text')}</p>")

        doc.close()
        return "\n".join(html_parts)

    def get_section_formatted(
        self,
        pdf_bytes: bytes,
        section: Section,
        output: str = "html",
        include_subsections: bool = True,
        strip_heading: bool = False,
    ) -> str | bytes:
        """Extract section content with formatting — no page chrome, just content.

        Reconstructs paragraph boundaries from Y-gaps between lines,
        detects lists, interleaves tables, and classifies heading levels.

        Args:
            section: The section to extract.
            output: "html" for semantic HTML, "docx" for Word document bytes.
            strip_heading: If True, remove the section's own heading from output.
                This gives the body content only — useful when the heading is
                already displayed by the UI (e.g., section tree, verbatim panel).
            include_subsections: If True, include subsection content.

        Returns:
            HTML string or DOCX bytes depending on output parameter.
        """
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        start = section.page
        declared_end = section.end_page if section.end_page is not None else start
        # Buffer: look up to 2 pages beyond declared end for boundary
        end = min(declared_end + 2, doc.page_count - 1)

        # Y-coordinate bounds
        start_y = section.start_y
        if start_y == 0.0:
            start_y = self._find_heading_y(doc, start, section.number, section.title)

        if not include_subsections and section.children:
            first_child = section.children[0]
            end = first_child.page
            end_y = self._find_heading_y(doc, end, first_child.number, first_child.title)
            if end_y == 0.0:
                end_y = 99999.0
        else:
            # Scan EACH page from declared_end onward for the next heading.
            # Use the resolved start_y (not 0.0) to correctly skip
            # the current section's own heading on a shared page.
            end_y = 99999.0
            for scan_page in range(declared_end, end + 1):
                after = start_y if scan_page == start else 0.0
                found_y = self._find_next_heading_y(
                    doc, scan_page, section.number, after_y=after,
                )
                if found_y < 99999.0:
                    end = scan_page
                    end_y = found_y
                    break

            # If no next heading found via generic scan, try finding
            # the specific next sibling section by number pattern.
            # Handles both same-page (2.2.1→2.2.2) and cross-page boundaries.
            if end_y >= 99999.0:
                next_num = self._next_section_number(section.number)
                if next_num:
                    for scan_page in range(start, end + 1):
                        sibling_y = self._find_heading_y(
                            doc, scan_page, next_num, ""
                        )
                        if sibling_y > 0:
                            # Same page: sibling must be below current heading
                            if scan_page == start and sibling_y <= start_y + 5:
                                continue
                            end = scan_page
                            end_y = sibling_y
                            break

        # Header/footer patterns to strip (static)
        hf_patterns = [
            re.compile(r"^Page \d+", re.IGNORECASE),
            re.compile(r"^Confidential$", re.IGNORECASE),
            re.compile(r"^\d{1,4}$"),
            re.compile(r"^PF-\d+", re.IGNORECASE),
            re.compile(r"^Protocol [A-Z]", re.IGNORECASE),
            re.compile(r"^Final Protocol", re.IGNORECASE),
            re.compile(r"^PFIZER CONFIDENTIAL$", re.IGNORECASE),
        ]

        # Finding 13: Auto-detect protocol-specific headers/footers by
        # finding text that appears in the top/bottom 80pt of 3+ pages
        from collections import Counter
        top_texts = Counter()
        bottom_texts = Counter()
        for page_num in range(start, min(end + 1, doc.page_count)):
            page = doc[page_num]
            page_height = page.rect.height
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if "lines" not in block:
                    continue
                for line in block["lines"]:
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    y = spans[0]["origin"][1]
                    text = "".join(s["text"] for s in spans).strip()
                    if not text or len(text) < 5:
                        continue
                    if y < 80:
                        top_texts[text] += 1
                    elif y > page_height - 80:
                        bottom_texts[text] += 1

        # Text appearing on 3+ pages in header/footer zone = auto-strip
        auto_hf: set[str] = set()
        pages_scanned = min(end + 1, doc.page_count) - start
        threshold = max(3, pages_scanned // 3)
        for text, count in top_texts.items():
            if count >= threshold:
                auto_hf.add(text)
        for text, count in bottom_texts.items():
            if count >= threshold:
                auto_hf.add(text)

        # Step 1: Extract all lines with full metadata
        raw_lines: list[dict] = []
        for page_num in range(start, end + 1):
            page = doc[page_num]
            page_start_y = start_y if page_num == start else 0.0
            page_end_y = end_y if page_num == end else 99999.0

            # Fix 2: Collect table bounding boxes to avoid duplicating
            # table content as flattened paragraphs
            table_bboxes = []
            try:
                found_tables = page.find_tables()
                for t in found_tables.tables:
                    if hasattr(t, "bbox"):
                        table_bboxes.append(t.bbox)  # (x0, y0, x1, y1)
            except (AttributeError, Exception):
                pass

            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if "lines" not in block:
                    continue
                for line in block["lines"]:
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    y = spans[0]["origin"][1]
                    if y < page_start_y - 2 or y > page_end_y - 2:
                        continue

                    # Skip text inside table bounding boxes — tables are
                    # extracted separately via find_tables() as structured data
                    if table_bboxes:
                        in_table = False
                        for bbox in table_bboxes:
                            if bbox[1] - 5 <= y <= bbox[3] + 5:
                                in_table = True
                                break
                        if in_table:
                            continue

                    # Join spans with space awareness — some PDFs have
                    # adjacent spans without whitespace between words
                    parts = []
                    for si, s in enumerate(spans):
                        t = s["text"]
                        if si > 0 and parts and t and parts[-1]:
                            # If previous span doesn't end with space and
                            # this span doesn't start with space, check if
                            # we need to insert one (word boundary)
                            prev = parts[-1]
                            if (not prev.endswith((" ", "-", "/"))
                                and not t.startswith((" ", ",", ".", ";", ":"))
                                and prev[-1].isalpha() and t[0].isalpha()):
                                parts.append(" ")
                        parts.append(t)
                    text = "".join(parts).strip()
                    if not text:
                        continue

                    # Skip headers/footers (static patterns + auto-detected)
                    if any(p.match(text) for p in hf_patterns):
                        continue
                    if text in auto_hf:
                        continue
                    if "Continued on table" in text:
                        continue

                    x = spans[0]["origin"][0]
                    font_size = max(s["size"] for s in spans)
                    is_bold = any(
                        "Bold" in s.get("font", "") or s["flags"] & 16
                        for s in spans if s["text"].strip()
                    )
                    is_italic = any(
                        "Italic" in s.get("font", "") or s["flags"] & 2
                        for s in spans if s["text"].strip()
                    )

                    raw_lines.append({
                        "text": text,
                        "y": y, "x": x,
                        "page": page_num,
                        "size": font_size,
                        "bold": is_bold,
                        "italic": is_italic,
                        "spans": spans,
                    })

            # Collect tables on this page within Y range
            try:
                found_tables = page.find_tables()
                for t in found_tables.tables:
                    if hasattr(t, "bbox"):
                        table_y = t.bbox[1]
                        if table_y < page_start_y - 10 or table_y > page_end_y + 10:
                            continue
                    raw_lines.append({
                        "text": None,
                        "y": t.bbox[1] if hasattr(t, "bbox") else y,
                        "x": 0, "page": page_num,
                        "size": 0, "bold": False, "italic": False,
                        "spans": [],
                        "table_data": t.extract(),
                    })
            except (AttributeError, Exception):
                pass

        doc.close()

        if not raw_lines:
            return "" if output == "html" else b""

        # Step 2: Paragraph reconstruction from Y-gaps
        paragraphs = self._reconstruct_paragraphs(raw_lines, section_number=section.number)

        # Step 3: Strip the section's own heading if requested
        if strip_heading and paragraphs:
            # Remove the first HEADING paragraph (it's the section title)
            if paragraphs[0].get("type") == "HEADING":
                paragraphs = paragraphs[1:]
            elif paragraphs[0].get("type") == "SUBHEADING":
                paragraphs = paragraphs[1:]

        # Step 4: Output generation
        if output == "docx":
            return self._paragraphs_to_docx(paragraphs)
        return self._paragraphs_to_html(paragraphs)

    def _reconstruct_paragraphs(self, raw_lines: list[dict], section_number: str = "") -> list[dict]:
        """Group consecutive lines into semantic paragraphs using Y-gaps.

        Grouping rules (from team review):
          Y-gap < 16pt AND same indent AND no bold change → same paragraph
          Y-gap >= 20pt OR indent change OR bold change → new paragraph
          Starts with "N." where N is digit, NOT bold → numbered list item
          X-indent >= 126 with bullet marker → nested list item
          Bold text followed by non-bold at same Y → split heading + body
        """
        if not raw_lines:
            return []

        raw_lines.sort(key=lambda l: (l["page"], l["y"]))

        paragraphs: list[dict] = []
        current: dict | None = None

        # Detect typical line spacing from same-page consecutive lines
        typical_gap = 14.0
        gaps = []
        for i in range(1, min(len(raw_lines), 30)):
            if raw_lines[i]["page"] == raw_lines[i - 1]["page"] and raw_lines[i].get("text"):
                gap = raw_lines[i]["y"] - raw_lines[i - 1]["y"]
                if 5 < gap < 20:
                    gaps.append(gap)
        if gaps:
            typical_gap = sorted(gaps)[len(gaps) // 2]  # Median, not mean

        # Threshold: Y-gap < 16pt = continuation, >= 20pt = new paragraph
        continuation_threshold = max(typical_gap + 2, 16.0)
        para_gap_threshold = max(typical_gap * 1.4, 20.0)

        # List item patterns — match ONLY at start of line
        # Finding 9: expanded bullet character set to catch Unicode variants
        numbered_list_re = re.compile(r"^(\d{1,3})[.)]\s")
        bullet_re = re.compile(
            r"^[\u2022\u2023\u25E6\u2043\u25AA\u25AB\u25CF\u25CB"
            r"\u2013\u2014\u2015\u2212\u25A0\u25A1\u25B8\u25B9"
            r"\uf0b7\uf0a7\uf0d8"  # Wingdings bullets from Word/PDF
            r"•●○◦◆◇■□▪▸►–—\-]\s"
        )

        # Compute base X (leftmost text position) for indent detection
        text_lines = [l for l in raw_lines if l.get("text")]
        base_x = min(l["x"] for l in text_lines) if text_lines else 0

        # Build set of known section numbers from multi-level headings in the
        # extracted lines. A line like "5.1. Inclusion Criteria" tells us that
        # "5" IS a real section. A line "4. Phase 2/3 only..." with no "4.1"
        # anywhere tells us "4" is NOT a section — it's a list item.
        known_section_nums: set[str] = set()
        known_parents: set[str] = set()
        for l in raw_lines:
            if l.get("text") and l.get("bold"):
                m = re.match(r"^(\d{1,2}(?:\.\d{1,2}){1,5})\.?\s+", l["text"])
                if m:
                    num = m.group(1)
                    known_section_nums.add(num)
                    # The parent of "5.1" is "5" — so "5" is a real section
                    parts = num.split(".")
                    if len(parts) >= 2:
                        known_parents.add(parts[0])
        known_section_nums |= known_parents

        # Indent levels (from Pfizer protocol analysis):
        # base_x (~72) = section body
        # base_x+18 (~90) = numbered criteria
        # base_x+36 (~108) = continuation / sub-paragraph
        # base_x+54 (~126) = bullet prefix position
        # base_x+72 (~144) = sub-bullet body text

        for line in raw_lines:
            # Table — always its own block
            if line.get("table_data") is not None:
                if current:
                    paragraphs.append(current)
                    current = None
                paragraphs.append({
                    "type": "TABLE",
                    "table_data": line["table_data"],
                    "y": line["y"],
                })
                continue

            text = line["text"]
            x = line["x"]
            indent_level = max(0, round((x - base_x) / 18))  # 18pt per indent

            # === CLASSIFY THIS LINE ===

            # Issue 2 fix: Section headings MUST be bold AND the number must
            # look like a real section number (not a list item like "4. Phase
            # 2/3 only..."). Heuristics:
            # - Single digits (1-9) at X > base_x+10 = list items, not sections
            # - Multi-level numbers (4.2, 5.1.3) = always section headings
            sec_match = re.match(r"^(\d{1,2}(?:\.\d{1,2}){0,5})\.?\s+[A-Z]", text)
            is_section_heading = False
            if line["bold"] and sec_match:
                candidate_num = sec_match.group(1)
                if "." in candidate_num:
                    # Multi-level number (4.2, 5.1.3) = section heading
                    is_section_heading = True
                elif section_number:
                    # We know which section we're extracting.
                    # Finding 10 fix: the section's OWN number is always a heading
                    # (it's the section title line like "1 Synopsis").
                    if candidate_num == section_number:
                        is_section_heading = True
                    # Subsection headings (5.1, 5.2 within section 5) are headings
                    elif candidate_num.startswith(section_number + "."):
                        is_section_heading = True
                    # Within Section 5.1, "4." is a list item, not section 4.
                    # A single digit is only a heading if it's a known section
                    # AND not the parent of our current section.
                    else:
                        sec_parent = section_number.split(".")[0] if "." in section_number else ""
                        if candidate_num in known_section_nums and candidate_num != sec_parent:
                            is_section_heading = True
                elif indent_level == 0:
                    is_section_heading = True

            # Issue 3/6 fix: Subheadings are bold standalone labels like
            # "Age and Sex:" or "Medical Conditions:". They are short, bold,
            # often end with ":", and are NOT list items or notes.
            # "Note: Healthy participants..." is a bold note, not a subheading.
            is_subheading = (
                line["bold"]
                and not is_section_heading
                and line["size"] >= 11
                and not numbered_list_re.match(text)
                and not text.lower().startswith("note")
                and len(text) < 80  # Subheadings are short
            )

            # Numbered list items: "1. Male or female..."
            # Can be bold (e.g., "4. Phase 2/3 only:") as long as they're
            # not classified as section headings above.
            is_numbered_list = bool(
                numbered_list_re.match(text) and not is_section_heading
            )

            # Bullet list items (• – ○ etc.)
            is_bullet = bool(bullet_re.match(text))

            is_list_item = is_numbered_list or is_bullet

            # === DECIDE: CONTINUE OR START NEW ===

            start_new = False
            if current is None:
                start_new = True

            # Issue 3/4 fix: Bold change = always new paragraph
            # "5.1. Inclusion Criteria" (bold) followed by "Participants are
            # eligible..." (not bold) must be separate paragraphs.
            elif line["bold"] != current.get("bold") and not (
                # Exception: don't split mid-line bold fragments like "Note: text"
                current.get("type") == "BODY" and not line["bold"]
                and line["page"] == current.get("last_page")
                and (line["y"] - current.get("last_y", 0)) < continuation_threshold
            ):
                start_new = True

            elif is_section_heading or is_subheading:
                start_new = True

            elif is_list_item and current.get("type") not in ("LIST_ITEM", "LIST_ITEM_L2"):
                start_new = True

            elif is_list_item and current.get("type") in ("LIST_ITEM", "LIST_ITEM_L2"):
                # New list item of same type = new paragraph
                # But continuation line of same item = merge
                if is_numbered_list or is_bullet:
                    start_new = True

            elif line["page"] != current.get("last_page"):
                # Cross-page boundary handling
                if is_section_heading or is_subheading:
                    start_new = True
                elif is_list_item:
                    # New list item on new page = always start new
                    start_new = True
                # Finding 8 fix: if current is a LIST_ITEM and the next line
                # is NOT a new list item, merge it as continuation.
                # "4. Phase 2/3 only: Participants who... are at higher risk"
                # continues on next page with "for acquiring COVID-19..."
                elif current and current["type"] in ("LIST_ITEM", "LIST_ITEM_L2") and not is_list_item:
                    # Check if previous text ends with sentence-ending punctuation
                    if current["text"].rstrip().endswith((".", ";", "!")):
                        start_new = True  # Sentence ended — new paragraph
                    # else: merge continuation into list item
                # Issue B: previous paragraph ends with period/colon AND this
                # line starts at a bullet-marker X position → new paragraph
                elif current and current["text"].rstrip().endswith((".", ":", ";")):
                    if is_bullet or indent_level >= 2:
                        start_new = True
                # Otherwise assume continuation (paragraph wraps across page)

            else:
                # Same page — check Y gap
                gap = line["y"] - current.get("last_y", 0)
                if gap >= para_gap_threshold:
                    start_new = True
                elif gap < continuation_threshold:
                    # Issue 1 fix: small gap + same indent + no bold change +
                    # no list marker = continuation of current paragraph.
                    # This merges wrapped lines like:
                    # "...ages of 18 and 55 years, inclusive, and 65 and"
                    # "85 years, inclusive (Phase 1)..."
                    pass  # Will merge below

            # Fix: merge heading fragments — when "4." and "Objectives and
            # Endpoints" are separate lines at same Y or tiny gap, join them
            # into one heading instead of creating separate paragraphs.
            if (
                start_new
                and is_section_heading
                and current
                and current["type"] == "HEADING"
                and current.get("last_page") == line["page"]
                and (line["y"] - current.get("last_y", 0)) < continuation_threshold
            ):
                # Merge into existing heading
                start_new = False

            # Fix: continuation lines after list items — if current is a
            # list item and this non-list-item line is indented, merge it
            # as continuation text of the list item.
            if (
                start_new
                and not is_list_item
                and not is_section_heading
                and not is_subheading
                and current
                and current["type"] in ("LIST_ITEM", "LIST_ITEM_L2")
                and current.get("last_page") == line["page"]
                and (line["y"] - current.get("last_y", 0)) < continuation_threshold
                and indent_level >= current.get("indent_level", 0)
            ):
                start_new = False

            if start_new:
                if current:
                    paragraphs.append(current)

                if is_section_heading:
                    ptype = "HEADING"
                elif is_subheading:
                    ptype = "SUBHEADING"
                elif is_bullet and indent_level >= 3:
                    ptype = "LIST_ITEM_L2"
                elif is_list_item:
                    ptype = "LIST_ITEM"
                else:
                    ptype = "BODY"

                current = {
                    "type": ptype,
                    "text": text,
                    "bold": line["bold"],
                    "italic": line["italic"],
                    "size": line["size"],
                    "x": x,
                    "indent_level": indent_level,
                    "y": line["y"],
                    "last_y": line["y"],
                    "last_page": line["page"],
                    "spans_data": [line["spans"]],
                    "is_numbered": is_numbered_list,
                }
            else:
                # Merge into current paragraph
                prev_text = current["text"]
                if prev_text.endswith("-"):
                    current["text"] = prev_text[:-1] + text
                elif prev_text.endswith(" ") or text.startswith(" "):
                    current["text"] += text
                else:
                    current["text"] += " " + text
                current["last_y"] = line["y"]
                current["last_page"] = line["page"]
                current["spans_data"].append(line["spans"])

        if current:
            paragraphs.append(current)

        return paragraphs

    def _paragraphs_to_html(self, paragraphs: list[dict]) -> str:
        """Convert reconstructed paragraphs to semantic HTML."""
        html_parts = []
        list_stack: list[str] = []  # Stack of open list tags for nesting

        def _close_lists_to(depth: int):
            while len(list_stack) > depth:
                html_parts.append(f"</{list_stack.pop()}>")

        for para in paragraphs:
            ptype = para["type"]

            # Close lists when leaving list context
            if ptype not in ("LIST_ITEM", "LIST_ITEM_L2"):
                _close_lists_to(0)

            if ptype == "TABLE":
                table_data = para.get("table_data", [])
                if table_data:
                    html_parts.append(self._table_to_html(table_data))

            elif ptype == "HEADING":
                text = self._escape_html(para["text"])
                match = re.match(r"^(\d{1,2}(?:\.\d{1,2}){0,5})", para["text"])
                level = len(match.group(1).split(".")) + 1 if match else 2
                level = min(level, 6)
                html_parts.append(f"<h{level}>{text}</h{level}>")

            elif ptype == "SUBHEADING":
                text = self._escape_html(para["text"])
                html_parts.append(f"<h4><strong>{text}</strong></h4>")

            elif ptype == "LIST_ITEM":
                # Close nested lists but keep L1 open if same type
                _close_lists_to(1)  # Close L2 nesting if any
                desired_type = "ol" if para.get("is_numbered") else "ul"
                if list_stack and list_stack[0] != desired_type:
                    # List type changed (ol→ul or ul→ol) — close and reopen
                    _close_lists_to(0)
                if not list_stack:
                    html_parts.append(f"<{desired_type}>")
                    list_stack.append(desired_type)
                item_text = re.sub(
                    r"^(?:[\u2022•●○\uf0b7\uf0a7]\s*|\d+[.)]\s*|[a-z][.)]\s*|[–—-]\s*)",
                    "", para["text"]
                )
                html_parts.append(f"  <li>{self._escape_html(item_text)}</li>")

            elif ptype == "LIST_ITEM_L2":
                # Ensure L1 list is open, then open L2
                if not list_stack:
                    html_parts.append("<ul>")
                    list_stack.append("ul")
                if len(list_stack) < 2:
                    html_parts.append("  <ul>")
                    list_stack.append("ul")
                item_text = re.sub(
                    r"^(?:[\u2022•●○\uf0b7\uf0a7]\s*|[–—-]\s*)",
                    "", para["text"]
                )
                html_parts.append(f"    <li>{self._escape_html(item_text)}</li>")

            else:  # BODY
                text = self._format_inline_html(para)
                html_parts.append(f"<p>{text}</p>")

        _close_lists_to(0)
        return "\n".join(html_parts)

    def _format_inline_html(self, para: dict) -> str:
        """Format paragraph text with inline bold/italic from span data."""
        # If we have span data, use it for precise formatting
        if para.get("spans_data"):
            parts = []
            for line_idx, spans in enumerate(para["spans_data"]):
                # Issue C fix: add space between merged lines to prevent
                # "orfemale" concatenation across line boundaries
                if line_idx > 0 and parts:
                    last = parts[-1]
                    if last and not last.endswith((" ", "-", ">")):
                        parts.append(" ")
                for span in spans:
                    text = self._escape_html(span.get("text", ""))
                    if not text.strip():
                        continue
                    flags = span.get("flags", 0)
                    is_bold = flags & 16 or "Bold" in span.get("font", "")
                    is_italic = flags & 2 or "Italic" in span.get("font", "")
                    if is_bold:
                        text = f"<strong>{text}</strong>"
                    if is_italic:
                        text = f"<em>{text}</em>"
                    parts.append(text)
            return "".join(parts)
        return self._escape_html(para["text"])

    @staticmethod
    def _escape_html(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    @staticmethod
    def _table_to_html(table_data: list[list]) -> str:
        """Convert table data to HTML table.

        Handles bullet characters inside table cells by converting them
        to proper HTML lists for clean rendering.
        """
        # Bullet chars commonly found in PDF table cells
        _bullet_chars = "\u25A0\u25AA\u25CF\u2022\u2023\u25E6\u2043\uf0b7\uf0a7•●○"
        _bullet_re = re.compile(rf"[{re.escape(_bullet_chars)}]\s*")

        def _format_cell(text: str) -> str:
            text = str(text or "").replace("<", "&lt;").replace(">", "&gt;")
            # If cell contains bullet chars, convert to HTML list
            if any(c in text for c in _bullet_chars):
                lines = _bullet_re.split(text)
                lines = [l.strip() for l in lines if l.strip()]
                if len(lines) > 1:
                    items = "".join(f"<li>{l}</li>" for l in lines)
                    return f"<ul>{items}</ul>"
            # If cell has newlines with content, split into paragraphs
            if "\n" in text:
                parts = [p.strip() for p in text.split("\n") if p.strip()]
                if len(parts) > 1:
                    return "<br>".join(parts)
            return text

        rows = []
        for i, row in enumerate(table_data):
            tag = "th" if i == 0 else "td"
            cells = "".join(
                f"<{tag}>{_format_cell(cell)}</{tag}>"
                for cell in row
            )
            rows.append(f"  <tr>{cells}</tr>")
        return "<table>\n" + "\n".join(rows) + "\n</table>"

    def _paragraphs_to_docx(self, paragraphs: list[dict]) -> bytes:
        """Convert reconstructed paragraphs to DOCX bytes."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        doc = Document()

        for para in paragraphs:
            ptype = para["type"]

            if ptype == "TABLE":
                table_data = para.get("table_data", [])
                if table_data and table_data[0]:
                    rows = len(table_data)
                    cols = max(len(r) for r in table_data)
                    tbl = doc.add_table(rows=rows, cols=cols, style="Table Grid")
                    for i, row in enumerate(table_data):
                        for j, cell in enumerate(row):
                            if j < cols:
                                tbl.cell(i, j).text = str(cell or "")

            elif ptype == "HEADING":
                match = re.match(r"^(\d{1,2}(?:\.\d{1,2}){0,5})", para["text"])
                level = len(match.group(1).split(".")) if match else 1
                level = min(level, 4)
                doc.add_heading(para["text"], level=level)

            elif ptype == "SUBHEADING":
                p = doc.add_paragraph()
                run = p.add_run(para["text"])
                run.bold = True

            elif ptype == "LIST_ITEM":
                item_text = re.sub(
                    r"^(?:[\u2022•●○]\s*|\d+[.)]\s*|[a-z][.)]\s*|[–—-]\s*)",
                    "", para["text"]
                )
                style = "List Number" if para.get("is_numbered") else "List Bullet"
                doc.add_paragraph(item_text, style=style)

            elif ptype == "LIST_ITEM_L2":
                item_text = re.sub(
                    r"^(?:[\u2022•●○]\s*|[–—-]\s*)",
                    "", para["text"]
                )
                doc.add_paragraph(item_text, style="List Bullet 2")

            else:  # BODY
                p = doc.add_paragraph()
                if para.get("spans_data"):
                    for spans in para["spans_data"]:
                        for span in spans:
                            text = span.get("text", "")
                            if not text:
                                continue
                            run = p.add_run(text)
                            flags = span.get("flags", 0)
                            run.bold = bool(flags & 16 or "Bold" in span.get("font", ""))
                            run.italic = bool(flags & 2 or "Italic" in span.get("font", ""))
                else:
                    p.add_run(para["text"])

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def get_tables_in_section(self, pdf_bytes: bytes, section: Section) -> list[dict]:
        """Extract tables from a specific section using PyMuPDF's table finder."""
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        start = section.page
        end = section.end_page if section.end_page is not None else start

        tables = []
        for page_num in range(start, min(end + 1, doc.page_count)):
            page = doc[page_num]
            # PyMuPDF 1.23+ has find_tables()
            try:
                found = page.find_tables()
                for t in found.tables:
                    table_data = t.extract()
                    if table_data:
                        tables.append({
                            "page": page_num,
                            "rows": len(table_data),
                            "cols": len(table_data[0]) if table_data else 0,
                            "data": table_data,
                        })
            except AttributeError:
                # Older PyMuPDF without find_tables
                pass

        doc.close()
        return tables

    def to_outline(self, sections: list[Section]) -> str:
        """Generate a readable outline of all sections."""
        lines = []
        for s in self._flatten(sections):
            indent = "  " * (s.level - 1)
            page_info = f"(p.{s.page_display})" if s.page_display else ""
            lines.append(f"{indent}{s.number} {s.title} {page_info}")
        return "\n".join(lines)

    # ------------------------------------------------------------------
    # Internal methods
    # ------------------------------------------------------------------

    def _parse_fitz_toc(self, toc: list) -> list[Section]:
        """Parse PyMuPDF's built-in TOC extraction."""
        sections = []
        for level, title, page in toc:
            # Try to extract section number from title
            match = re.match(r"^(\d{1,2}(?:\.\d{1,3}){0,4})\.?\s+(.*)", title.strip())
            if match:
                sec_num = match.group(1)
                # Finding 6 fix: compute level from section number instead
                # of trusting fitz TOC level. P-14's PDF bookmarks assign
                # level 1 to subsections like "2.1", "4.1" which should be L2.
                computed_level = sec_num.count(".") + 1
                sections.append(Section(
                    number=sec_num,
                    title=match.group(2).strip(),
                    page=page - 1,  # fitz TOC is 1-indexed
                    page_display=str(page),
                    level=computed_level,
                ))
            else:
                # Non-numbered sections (appendices, etc.)
                sections.append(Section(
                    number="",
                    title=title.strip(),
                    page=page - 1,
                    page_display=str(page),
                    level=level,
                ))
        return sections

    def _parse_toc_text(self, doc: fitz.Document) -> list[Section]:
        """Find and parse the Table of Contents from page text."""
        sections = []

        # Look for TOC — scan up to first 30 pages (some protocols have
        # long synopses before the TOC)
        for page_num in range(min(30, doc.page_count)):
            page = doc[page_num]
            text = page.get_text("text")

            # Check if this page looks like a TOC
            has_toc_header = any(kw in text.upper() for kw in
                                ["TABLE OF CONTENTS", "CONTENTS", "LIST OF TABLES"])
            # Also check if this page has dotted TOC lines (continuation pages)
            has_dotted_lines = text.count("...") > 3 or text.count("…") > 3 or text.count("..") > 5
            if not has_toc_header and not has_dotted_lines:
                if page_num > 10:
                    continue

            for line in text.split("\n"):
                line = line.strip()
                if not line:
                    continue

                # Try dotted TOC format
                match = _TOC_RE.match(line)
                if not match:
                    match = _TOC_SIMPLE_RE.match(line)
                if match:
                    sec_num = match.group(1)
                    sec_title = match.group(2).strip()
                    sec_page = match.group(3)
                    level = sec_num.count(".") + 1
                    sections.append(Section(
                        number=sec_num,
                        title=sec_title,
                        page=int(sec_page) - 1,  # Will be calibrated below
                        page_display=sec_page,
                        level=level,
                    ))

        # Calibrate page offset: the TOC page numbers may not match
        # PDF physical page numbers (cover pages, roman numeral preface, etc.)
        if sections:
            sections = self._calibrate_page_offset(doc, sections)

        return sections

    def _calibrate_page_offset(self, doc: fitz.Document, sections: list[Section]) -> list[Section]:
        """Calibrate TOC page numbers by finding each section header in actual PDF text.

        The most reliable approach: for each top-level section, scan the PDF
        to find the page where that header actually appears.
        """
        # Build a quick page→text index for the document
        page_texts: dict[int, str] = {}
        for i in range(doc.page_count):
            page_texts[i] = doc[i].get_text("text")

        # First, detect offset from Section 1 (or first numbered section)
        offset = 0
        for s in sections:
            if s.number and re.match(r"^\d$", s.number) and len(s.title) > 3:
                title_word = re.sub(r"[^A-Za-z]", "", s.title[:15]).strip()[:10]
                if not title_word:
                    continue
                # Search AFTER the TOC pages (skip pages where TOC was found)
                # TOC is typically in the first 30 pages; start after s.page
                search_from = max(s.page + 1, 5)  # At least page 5, after TOC
                for page_num in range(search_from, doc.page_count):
                    pattern = r"(?:^|\n)\s*" + re.escape(s.number) + r"[\.\s]+" + title_word
                    if re.search(pattern, page_texts[page_num], re.IGNORECASE):
                        offset = page_num - s.page
                        logger.info(f"Page offset: Section {s.number} '{s.title[:20]}' "
                                    f"TOC={s.page} actual={page_num} offset={offset:+d}")
                        break
                break

        if offset != 0:
            for s in sections:
                s.page = max(0, s.page + offset)

        return sections

    def _scan_headers(self, doc: fitz.Document) -> list[Section]:
        """Scan all pages for section headers using text + font analysis.

        Uses two strategies:
        1. Regex matching on section number patterns (with date rejection)
        2. Font-size detection for bold/large text that may be unnumbered headings
        """
        sections = []
        seen = set()

        # Month names for date rejection
        months = {"january", "february", "march", "april", "may", "june",
                  "july", "august", "september", "october", "november", "december",
                  "jan", "feb", "mar", "apr", "jun", "jul", "aug", "sep", "oct", "nov", "dec"}

        # ICH standard section titles for validation
        standard_titles = {
            # Clinical protocol sections
            "introduction", "background", "objectives", "endpoints",
            "study design", "study population", "inclusion", "exclusion",
            "study intervention", "discontinuation", "assessments",
            "procedures", "statistical", "considerations", "references",
            "appendix", "supporting", "documentation", "abbreviations",
            "synopsis", "summary", "schedule of activities",
            "adverse events", "safety", "efficacy", "laboratory",
            "pharmacokinetic", "pharmacodynamic", "dosing",
            # General document / contract sections
            "definitions", "scope", "purpose", "terms", "conditions",
            "obligations", "responsibilities", "payment", "compensation",
            "indemnification", "indemnity", "confidentiality", "termination",
            "amendment", "governing law", "jurisdiction", "dispute",
            "insurance", "liability", "warranty", "representations",
            "compliance", "regulatory", "data protection", "privacy",
            "intellectual property", "publication", "force majeure",
            "notices", "general provisions", "miscellaneous", "signatures",
            "exhibits", "schedules", "attachments", "budget", "fees",
            "milestones", "deliverables", "timeline", "protocol",
            "investigator", "sponsor", "site", "institution",
            "clinical trial agreement", "overview", "recitals",
        }

        for page_num in range(doc.page_count):
            page = doc[page_num]

            # Strategy 1: Regex on plain text with date rejection
            text = page.get_text("text")
            for line in text.split("\n"):
                line = line.strip()
                # Try numbered, lettered, and Roman numeral patterns
                match = _SECTION_RE.match(line)
                if not match:
                    match = _SECTION_LETTER_RE.match(line)
                if not match:
                    match = _SECTION_ROMAN_RE.match(line)
                if match:
                    sec_num = match.group(1)
                    sec_title = match.group(2).strip()

                    # REJECT DATES: "7 June", "10 October", "29 December"
                    if sec_title and sec_title.split() and sec_title.lower().split()[0] in months:
                        continue
                    # Reject very short "titles" that are likely not sections
                    if len(sec_title) < 3:
                        continue
                    # Reject if title is just a number or date-like
                    if re.match(r"^\d", sec_title) and not any(
                        kw in sec_title.lower() for kw in standard_titles
                    ):
                        continue

                    key = f"{sec_num}_{sec_title[:20]}"
                    if key in seen:
                        continue
                    seen.add(key)

                    level = sec_num.count(".") + 1
                    sections.append(Section(
                        number=sec_num,
                        title=sec_title,
                        page=page_num,
                        page_display=str(page_num + 1),
                        level=level,
                    ))

            # Strategy 2: Font-size detection for bold headers without numbers
            # This catches headers that are just bold text (no section number)
            try:
                blocks = page.get_text("dict")["blocks"]
                for block in blocks:
                    if block.get("type") != 0:  # text blocks only
                        continue
                    for line_data in block.get("lines", []):
                        for span in line_data.get("spans", []):
                            text_span = span.get("text", "").strip()
                            font_size = span.get("size", 0)
                            flags = span.get("flags", 0)
                            is_bold = flags & 2 ** 4  # bit 4 = bold

                            # Large bold text (>12pt) that looks like a heading
                            if font_size >= 12 and is_bold and len(text_span) > 3:
                                text_lower = text_span.lower()
                                # Accept if: has a known keyword OR is uppercase OR
                                # looks like a title (capitalized words, no period at end)
                                is_title_case = text_span[0].isupper() and not text_span.endswith(".")
                                is_all_upper = text_span == text_span.upper() and len(text_span) > 3
                                has_keyword = any(kw in text_lower for kw in standard_titles)

                                if has_keyword or is_all_upper or (is_title_case and len(text_span) > 5):
                                    # Try to extract section number
                                    num_match = re.match(
                                        r"^(\d{1,2}(?:\.\d{1,3}){0,4})\.?\s+(.*)",
                                        text_span
                                    )
                                    if num_match:
                                        sec_num = num_match.group(1)
                                        sec_title = num_match.group(2).strip()
                                    else:
                                        sec_num = ""
                                        sec_title = text_span

                                    key = f"font_{sec_num}_{sec_title[:20]}"
                                    if key in seen:
                                        continue
                                    seen.add(key)

                                    level = sec_num.count(".") + 1 if sec_num else 1
                                    sections.append(Section(
                                        number=sec_num,
                                        title=sec_title,
                                        page=page_num,
                                        page_display=str(page_num + 1),
                                        level=level,
                                    ))
            except Exception:
                pass  # Font analysis is best-effort

        return sections

    def _compute_page_ranges(self, sections: list[Section], total_pages: int):
        """Compute end_page for each section and build parent-child hierarchy."""
        flat = self._flatten(sections)

        # Deduplicate: same section number + page should appear only once.
        # Duplicates cause wrong end_page (next entry is yourself, not the next section).
        seen: set[tuple[str, int]] = set()
        deduped: list[Section] = []
        for s in flat:
            key = (s.number, s.page)
            if key in seen:
                continue
            seen.add(key)
            deduped.append(s)
        flat = deduped

        for i, s in enumerate(flat):
            if i + 1 < len(flat):
                s.end_page = flat[i + 1].page - 1
                if s.end_page < s.page:
                    s.end_page = s.page
            else:
                s.end_page = total_pages - 1

        # Finding 11: Build children tree from section numbers
        # "2.1" is a child of "2", "4.1.1" is a child of "4.1"
        section_map = {s.number: s for s in flat if s.number}
        for s in flat:
            if not s.number or "." not in s.number:
                continue
            parent_num = ".".join(s.number.split(".")[:-1])
            parent = section_map.get(parent_num)
            if parent and s not in parent.children:
                parent.children.append(s)

    def _flatten(self, sections: list[Section]) -> list[Section]:
        """Flatten nested sections into a sorted list."""
        result = []
        for s in sections:
            result.append(s)
            result.extend(self._flatten(s.children))
        return sorted(result, key=lambda s: (s.page, s.number))
