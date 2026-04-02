"""
DOCX Ingestor — converts DOCX documents to FormattedDocument IR.

Reads paragraph formatting (bold, italic, underline, font, size, color),
tables with merged cells, heading styles, and list styles from Word documents.

Handles style inheritance: checks run-level, then paragraph style, then
document defaults for formatting attributes.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Any

from src.formatter.extractor import (
    FormattedDocument,
    FormattedLine,
    FormattedPage,
    FormattedParagraph,
    FormattedSpan,
    FormattedTable,
    FormattedTableCell,
)

logger = logging.getLogger(__name__)

_DEFAULT_FONT = "Calibri"
_DEFAULT_SIZE = 11.0


class DOCXIngestor:
    """Converts DOCX bytes into a FormattedDocument IR.

    Walks the style inheritance chain: run → paragraph style → document
    defaults for each formatting attribute.
    """

    def ingest(self, content: bytes, filename: str = "") -> FormattedDocument:
        """Parse a DOCX file and return a FormattedDocument."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            raise ImportError("python-docx is required: pip install python-docx")

        doc = Document(io.BytesIO(content))

        # Extract document default styles
        default_font, default_size = self._get_doc_defaults(doc)

        pages: list[FormattedPage] = []
        paragraphs: list[FormattedParagraph] = []
        tables: list[FormattedTable] = []
        font_counts: dict[str, int] = {}
        color_counts: dict[str, int] = {}
        style_counts: dict[str, int] = {}

        # Process paragraphs
        for para in doc.paragraphs:
            fmt_para = self._process_paragraph(
                para, default_font, default_size, font_counts, color_counts,
            )
            if fmt_para:
                paragraphs.append(fmt_para)
                style_counts[fmt_para.style] = style_counts.get(fmt_para.style, 0) + 1

        # Process tables
        for table in doc.tables:
            fmt_table = self._process_table(table, font_counts)
            if fmt_table:
                tables.append(fmt_table)

        # Create single page (DOCX doesn't have explicit pages at parse time)
        page = FormattedPage(
            page_number=0,
            width=595.3,   # A4 default
            height=842.0,
            paragraphs=paragraphs,
            tables=tables,
        )

        # Try to get actual page dimensions from section
        if doc.sections:
            section = doc.sections[0]
            try:
                page.width = section.page_width.pt
                page.height = section.page_height.pt
                page.margin_left = section.left_margin.pt
                page.margin_right = section.right_margin.pt
                page.margin_top = section.top_margin.pt
                page.margin_bottom = section.bottom_margin.pt
            except Exception:
                pass

        pages.append(page)

        return FormattedDocument(
            filename=filename,
            pages=pages,
            font_inventory=font_counts,
            color_inventory=color_counts,
            style_inventory=style_counts,
        )

    def _get_doc_defaults(self, doc: Any) -> tuple[str, float]:
        """Extract document-level default font and size."""
        font = _DEFAULT_FONT
        size = _DEFAULT_SIZE

        try:
            # Try to get from default paragraph style
            default_style = doc.styles["Normal"]
            if default_style.font.name:
                font = default_style.font.name
            if default_style.font.size:
                size = default_style.font.size.pt
        except Exception:
            pass

        return font, size

    def _resolve_bold(self, run: Any, para_style: Any) -> bool:
        """Resolve bold through inheritance chain: run → paragraph style → defaults."""
        # Explicit run-level bold
        if run.font.bold is True:
            return True
        if run.font.bold is False:
            return False

        # Run bold is None — check paragraph style
        try:
            if para_style and para_style.font.bold:
                return True
        except Exception:
            pass

        return False

    def _resolve_italic(self, run: Any, para_style: Any) -> bool:
        """Resolve italic through inheritance chain."""
        if run.font.italic is True:
            return True
        if run.font.italic is False:
            return False

        try:
            if para_style and para_style.font.italic:
                return True
        except Exception:
            pass

        return False

    def _resolve_font_name(self, run: Any, para_style: Any, default_font: str) -> str:
        """Resolve font name through inheritance chain."""
        if run.font.name:
            return run.font.name

        try:
            if para_style and para_style.font.name:
                return para_style.font.name
        except Exception:
            pass

        return default_font

    def _resolve_font_size(self, run: Any, para_style: Any, default_size: float) -> float:
        """Resolve font size through inheritance chain."""
        if run.font.size:
            return run.font.size.pt

        try:
            if para_style and para_style.font.size:
                return para_style.font.size.pt
        except Exception:
            pass

        return default_size

    def _resolve_color(self, run: Any) -> int:
        """Extract color from run, return as packed RGB integer."""
        try:
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                return (int(rgb[0:2], 16) << 16) | (int(rgb[2:4], 16) << 8) | int(rgb[4:6], 16)
        except Exception:
            pass
        return 0

    def _process_paragraph(
        self,
        para: Any,
        default_font: str,
        default_size: float,
        font_counts: dict[str, int],
        color_counts: dict[str, int],
    ) -> FormattedParagraph | None:
        """Convert a python-docx Paragraph to FormattedParagraph."""
        if not para.runs and not para.text.strip():
            return None

        # Get paragraph style for inheritance
        para_style = None
        try:
            para_style = para.style
        except Exception:
            pass

        # Determine paragraph style classification
        style_name = para.style.name if para.style else "Normal"
        fmt_style = "body"

        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
                fmt_style = f"heading{min(level, 6)}"
            except (ValueError, IndexError):
                fmt_style = "heading3"
        elif "List Bullet" in style_name:
            fmt_style = "list_bullet"
        elif "List Number" in style_name:
            fmt_style = "list_number"
        elif "Title" in style_name:
            fmt_style = "heading1"
        elif "Subtitle" in style_name:
            fmt_style = "heading2"

        # Process runs
        spans: list[FormattedSpan] = []
        x_pos = 0.0

        if para.runs:
            for run in para.runs:
                text = run.text
                if not text:
                    continue

                font_name = self._resolve_font_name(run, para_style, default_font)
                font_size = self._resolve_font_size(run, para_style, default_size)
                is_bold = self._resolve_bold(run, para_style)
                is_italic = self._resolve_italic(run, para_style)
                is_underline = run.font.underline or False
                is_superscript = run.font.superscript or False
                is_subscript = run.font.subscript or False
                color = self._resolve_color(run)

                char_width = font_size * 0.5
                span_width = len(text) * char_width

                span = FormattedSpan(
                    text=text,
                    x0=x_pos,
                    y0=0,
                    x1=x_pos + span_width,
                    y1=font_size,
                    font=font_name,
                    size=font_size,
                    color=color,
                    bold=is_bold,
                    italic=is_italic,
                    underline=is_underline,
                    superscript=is_superscript,
                    subscript=is_subscript,
                )
                spans.append(span)
                x_pos += span_width

                # Track inventories
                font_counts[font_name] = font_counts.get(font_name, 0) + 1
                hex_color = f"#{color:06X}"
                color_counts[hex_color] = color_counts.get(hex_color, 0) + 1

        elif para.text.strip():
            # Paragraph with text but no runs (rare)
            text = para.text
            spans.append(FormattedSpan(
                text=text,
                x0=0, y0=0,
                x1=len(text) * default_size * 0.5,
                y1=default_size,
                font=default_font,
                size=default_size,
            ))

        if not spans:
            return None

        # Detect indent level
        indent_level = 0
        try:
            pf = para.paragraph_format
            if pf.left_indent:
                indent_level = max(0, int(pf.left_indent.pt / 18))
        except Exception:
            pass

        # Detect alignment
        alignment = "left"
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            align_map = {
                WD_ALIGN_PARAGRAPH.CENTER: "center",
                WD_ALIGN_PARAGRAPH.RIGHT: "right",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
            }
            if para.alignment in align_map:
                alignment = align_map[para.alignment]
        except Exception:
            pass

        line = FormattedLine(spans=spans, y_center=0, indent=x_pos * 0 if not spans else spans[0].x0)

        return FormattedParagraph(
            lines=[line],
            style=fmt_style,
            indent_level=indent_level,
            alignment=alignment,
        )

    def _process_table(
        self, table: Any, font_counts: dict[str, int],
    ) -> FormattedTable | None:
        """Convert a python-docx Table to FormattedTable."""
        rows = []
        num_rows = len(table.rows)
        num_cols = len(table.columns) if table.rows else 0

        for r_idx, row in enumerate(table.rows):
            fmt_row: list[FormattedTableCell] = []
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()

                # Detect bold (check first run of first paragraph)
                is_bold = False
                try:
                    first_para = cell.paragraphs[0] if cell.paragraphs else None
                    if first_para and first_para.runs:
                        is_bold = first_para.runs[0].font.bold or False
                except Exception:
                    pass

                # Detect merged cells via grid span
                colspan = 1
                rowspan = 1
                try:
                    tc = cell._tc
                    grid_span = tc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
                    if grid_span:
                        colspan = int(grid_span)
                    v_merge = tc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
                    if v_merge is not None:
                        # vMerge="restart" starts a merge, vMerge="" continues
                        if v_merge == "restart":
                            rowspan = 1  # will be computed later
                except Exception:
                    pass

                fmt_row.append(FormattedTableCell(
                    text=text,
                    row=r_idx,
                    col=c_idx,
                    rowspan=rowspan,
                    colspan=colspan,
                    bold=is_bold,
                    is_header=(r_idx == 0),
                ))

                # Track font
                font_counts[_DEFAULT_FONT] = font_counts.get(_DEFAULT_FONT, 0) + 1

            rows.append(fmt_row)

        if not rows:
            return None

        return FormattedTable(
            rows=rows,
            num_rows=num_rows,
            num_cols=num_cols,
        )
