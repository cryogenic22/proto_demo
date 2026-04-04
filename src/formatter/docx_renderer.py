"""
DOCX Renderer — converts FormattedDocument IR to a styled Word document.

Preserves: fonts, sizes, colors, bold/italic/underline, super/subscript,
paragraph spacing, alignment, heading levels, lists, table structure,
page margins, page breaks, and formula sub/superscript formatting.

Usage:
    from src.formatter.docx_renderer import DOCXRenderer
    renderer = DOCXRenderer()
    docx_bytes = renderer.render(formatted_document)
    # or render from PDF directly:
    docx_bytes = renderer.render_from_pdf(pdf_bytes)
"""

from __future__ import annotations

import io
import logging
import re
from typing import Any

from src.formatter.extractor import (
    FormattedDocument,
    FormattedPage,
    FormattedParagraph,
    FormattedSpan,
    FormattingExtractor,
)

logger = logging.getLogger(__name__)

# Default fallback font when source font is empty or unresolvable
_FALLBACK_FONT = "Calibri"

# Regex to split formula HTML into normal text and <sub>/<sup> segments
_FORMULA_TAG_RE = re.compile(r"(<su[bp]>)(.*?)(</su[bp]>)", re.DOTALL)


def _strip_control_chars(text: str) -> str:
    """Strip XML-invalid control characters while preserving tab/newline/return.

    XML 1.0 allows: #x9 (tab), #xA (LF), #xD (CR), #x20-#xD7FF, etc.
    We strip 0x00-0x08, 0x0B, 0x0C, 0x0E-0x1F which are all invalid.
    """
    return "".join(
        c for c in text
        if c in ("\t", "\n", "\r") or ord(c) >= 0x20
    )


def _parse_formula_html(html: str) -> list[tuple[str, str]]:
    """Parse formula HTML into a list of (text, tag_type) tuples.

    tag_type is one of: "normal", "sub", "sup".

    Example:
        "AUC<sub>0-inf</sub>" -> [("AUC", "normal"), ("0-inf", "sub")]
        "CO<sub>2</sub>" -> [("CO", "normal"), ("2", "sub")]
        "10<sup>6</sup> cells" -> [("10", "normal"), ("6", "sup"), (" cells", "normal")]
    """
    segments: list[tuple[str, str]] = []
    pos = 0

    for m in _FORMULA_TAG_RE.finditer(html):
        # Text before this tag
        if m.start() > pos:
            pre = html[pos:m.start()]
            if pre:
                segments.append((pre, "normal"))
        # The tagged content
        open_tag = m.group(1)
        content = m.group(2)
        tag_type = "sub" if "sub" in open_tag else "sup"
        if content:
            segments.append((content, tag_type))
        pos = m.end()

    # Remaining text after last tag
    if pos < len(html):
        remaining = html[pos:]
        if remaining:
            segments.append((remaining, "normal"))

    return segments


def _set_cell_borders(cell_element: Any, border_color: str = "000000", border_size: str = "4") -> None:
    """Set all four borders on a table cell via XML.

    Args:
        cell_element: The w:tc OxmlElement.
        border_color: Hex color string (e.g. "000000" for black).
        border_size: Border width in eighths of a point (4 = 0.5pt).
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc_pr = cell_element.get_or_add_tcPr()

    # Remove existing borders if any
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)

    borders_el = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        edge_el = OxmlElement(f"w:{edge}")
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), border_size)
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), border_color)
        borders_el.append(edge_el)
    tc_pr.append(borders_el)


def _set_cell_margin(cell_element: Any, margin_twips: int = 72) -> None:
    """Set uniform cell padding (margins) on a table cell.

    Args:
        cell_element: The w:tc OxmlElement.
        margin_twips: Margin size in twips (72 twips ~ 1mm, 144 ~ 2mm).
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc_pr = cell_element.get_or_add_tcPr()

    # Remove existing margins
    existing = tc_pr.find(qn("w:tcMar"))
    if existing is not None:
        tc_pr.remove(existing)

    mar_el = OxmlElement("w:tcMar")
    for edge in ("top", "left", "bottom", "right"):
        edge_el = OxmlElement(f"w:{edge}")
        edge_el.set(qn("w:w"), str(margin_twips))
        edge_el.set(qn("w:type"), "dxa")
        mar_el.append(edge_el)
    tc_pr.append(mar_el)


class DOCXRenderer:
    """Renders a FormattedDocument to DOCX with full formatting preservation."""

    def __init__(self):
        self.extractor = FormattingExtractor()

    def render(self, doc: FormattedDocument) -> bytes:
        """Render a FormattedDocument to DOCX bytes."""
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        word_doc = Document()

        # Set page margins from first page
        if doc.pages:
            p0 = doc.pages[0]
            for section in word_doc.sections:
                # Convert PDF points to EMU (1 pt = 12700 EMU)
                section.left_margin = Emu(int(p0.margin_left * 12700))
                section.right_margin = Emu(int(p0.margin_right * 12700))
                section.top_margin = Emu(int(p0.margin_top * 12700))
                section.bottom_margin = Emu(int(p0.margin_bottom * 12700))
                # Set page size
                section.page_width = Emu(int(p0.width * 12700))
                section.page_height = Emu(int(p0.height * 12700))

        for page_idx, page in enumerate(doc.pages):
            # Insert page break before every page except the first
            if page_idx > 0:
                self._add_page_break(word_doc)

            # Interleave paragraphs and tables by Y position
            for para in page.paragraphs:
                # Skip header/footer paragraphs — they belong in Word
                # headers/footers, not the document body
                if para.style in ("header", "footer"):
                    continue
                if para.style == "image":
                    self._add_image(word_doc, para)
                else:
                    self._add_paragraph(word_doc, para)

            # Add tables
            for table in page.tables:
                if not table.is_empty:
                    self._add_table(word_doc, table)

        # Write to bytes
        buf = io.BytesIO()
        word_doc.save(buf)
        return buf.getvalue()

    def render_from_pdf(self, pdf_bytes: bytes, filename: str = "") -> bytes:
        """Extract formatting from PDF and render to DOCX."""
        doc = self.extractor.extract(pdf_bytes, filename)
        return self.render(doc)

    def render_with_profile(
        self, doc: FormattedDocument, profile: dict[str, Any] | None = None,
    ) -> bytes:
        """Render with an optional style profile override.

        If a profile is provided (from TemplateConformer), use template
        fonts/sizes instead of source document formatting.
        """
        from docx import Document
        from docx.shared import Pt, RGBColor, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        word_doc = Document()

        # Set page margins
        if doc.pages:
            p0 = doc.pages[0]
            if profile:
                ml = profile.get("margin_left", p0.margin_left)
                mr = profile.get("margin_right", p0.margin_right)
                mt = profile.get("margin_top", p0.margin_top)
                mb = profile.get("margin_bottom", p0.margin_bottom)
            else:
                ml, mr, mt, mb = p0.margin_left, p0.margin_right, p0.margin_top, p0.margin_bottom

            for section in word_doc.sections:
                section.left_margin = Emu(int(ml * 12700))
                section.right_margin = Emu(int(mr * 12700))
                section.top_margin = Emu(int(mt * 12700))
                section.bottom_margin = Emu(int(mb * 12700))
                section.page_width = Emu(int(p0.width * 12700))
                section.page_height = Emu(int(p0.height * 12700))

        for page_idx, page in enumerate(doc.pages):
            # Page breaks between pages
            if page_idx > 0:
                self._add_page_break(word_doc)

            for para in page.paragraphs:
                if para.style in ("header", "footer"):
                    continue
                self._add_paragraph(word_doc, para, profile)

        buf = io.BytesIO()
        word_doc.save(buf)
        return buf.getvalue()

    def _add_page_break(self, word_doc: Any) -> None:
        """Insert a page break into the document."""
        from docx.enum.text import WD_BREAK
        p = word_doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

    def _add_paragraph(
        self,
        word_doc: Any,
        para: FormattedParagraph,
        profile: dict[str, Any] | None = None,
    ) -> None:
        """Add a formatted paragraph to the Word document."""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Determine Word style — heading1 through heading6 map to built-in
        # "Heading 1" through "Heading 6" so they appear in navigation pane
        style_name = None
        if para.style.startswith("heading"):
            # Extract level: "heading1" -> "1", "heading12" -> "1" (take last digit)
            digits = "".join(c for c in para.style if c.isdigit())
            level = int(digits) if digits else 3
            level = max(1, min(level, 6))  # Clamp to 1-6
            style_name = f"Heading {level}"
        elif para.style == "list_bullet":
            style_name = "List Bullet"
        elif para.style == "list_number":
            style_name = "List Number"

        try:
            wp = word_doc.add_paragraph(style=style_name)
        except KeyError:
            wp = word_doc.add_paragraph()

        # Set alignment
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        wp.alignment = align_map.get(para.alignment, WD_ALIGN_PARAGRAPH.LEFT)

        # Set spacing
        pf = wp.paragraph_format
        if para.spacing_before > 0:
            pf.space_before = Pt(min(para.spacing_before, 36))
        if para.spacing_after > 0:
            pf.space_after = Pt(min(para.spacing_after, 36))

        # Set indent
        if para.indent_level > 0 and not style_name:
            pf.left_indent = Pt(para.indent_level * 18)

        # Add runs with formatting
        for line in para.lines:
            for span in line.spans:
                if not span.text:
                    continue

                # Check if this span has a formula with HTML sub/sup tags
                # that we can render as separate Word runs
                if span.formula and self._render_formula_runs(wp, span, para, profile):
                    continue

                # Standard span rendering (no formula or formula without HTML)
                self._add_span_run(wp, span, para, profile)

    def _render_formula_runs(
        self,
        wp: Any,
        span: FormattedSpan,
        para: FormattedParagraph,
        profile: dict[str, Any] | None = None,
    ) -> bool:
        """Attempt to render a formula span as multiple Word runs with sub/sup.

        Returns True if formula HTML was successfully parsed and rendered,
        False if we should fall back to the standard span rendering.
        """
        from docx.shared import Pt, RGBColor

        formula = span.formula
        # Get the best HTML representation
        html = ""
        if hasattr(formula, "html") and formula.html:
            html = formula.html
        elif hasattr(formula, "plain_text") and formula.plain_text:
            # No HTML available — fall back to standard rendering
            return False

        if not html or ("<sub>" not in html and "<sup>" not in html):
            return False

        segments = _parse_formula_html(html)
        if not segments:
            return False

        # Resolve font and size for this span
        font_name = self._resolve_font(span, para, profile)
        size = self._resolve_size(span, para, profile)

        for seg_text, seg_type in segments:
            seg_text = _strip_control_chars(seg_text)
            if not seg_text:
                continue

            run = wp.add_run(seg_text)
            run.font.name = font_name
            run.font.size = Pt(size)
            run.font.bold = span.bold
            run.font.italic = span.italic
            if span.underline:
                run.font.underline = True

            # Apply sub/sup from the formula HTML tag
            if seg_type == "sub":
                run.font.subscript = True
            elif seg_type == "sup":
                run.font.superscript = True
            else:
                # Normal segment — respect span-level sub/sup flags
                if span.superscript:
                    run.font.superscript = True
                elif span.subscript:
                    run.font.subscript = True

            # Color
            r, g, b = span.color_rgb
            if r > 20 or g > 20 or b > 20:
                run.font.color.rgb = RGBColor(r, g, b)

        return True

    def _add_span_run(
        self,
        wp: Any,
        span: FormattedSpan,
        para: FormattedParagraph,
        profile: dict[str, Any] | None = None,
    ) -> None:
        """Add a single Word run for a span (standard path, no formula splitting)."""
        from docx.shared import Pt, RGBColor

        # Ensure text is a valid string (guard against bytes/None)
        text = str(span.text) if span.text else ""
        text = _strip_control_chars(text)
        if not text:
            return

        run = wp.add_run(text)

        # Font name (with fallback)
        font_name = self._resolve_font(span, para, profile)
        run.font.name = font_name

        # Font size
        size = self._resolve_size(span, para, profile)
        run.font.size = Pt(size)

        # Bold / Italic
        run.font.bold = span.bold
        run.font.italic = span.italic

        # Underline
        if span.underline:
            run.font.underline = True

        # Superscript / Subscript
        if span.superscript:
            run.font.superscript = True
        elif span.subscript:
            run.font.subscript = True

        # Color — preserve ALL non-black colors (red instructions,
        # blue links, grey headers)
        r, g, b = span.color_rgb
        if r > 20 or g > 20 or b > 20:
            run.font.color.rgb = RGBColor(r, g, b)

    def _resolve_font(
        self, span: FormattedSpan, para: FormattedParagraph,
        profile: dict[str, Any] | None = None,
    ) -> str:
        """Resolve font name for a span, applying profile overrides and fallback."""
        font_name = span.font_family
        if profile and not para.style.startswith("heading"):
            font_name = profile.get("body_font", font_name)
        elif profile and para.style.startswith("heading"):
            heading_fonts = profile.get("heading_fonts", {})
            font_name = heading_fonts.get(para.style, font_name)
        # Fallback to Calibri if font is empty or blank
        if not font_name or not font_name.strip():
            font_name = _FALLBACK_FONT
        return font_name

    def _resolve_size(
        self, span: FormattedSpan, para: FormattedParagraph,
        profile: dict[str, Any] | None = None,
    ) -> float:
        """Resolve font size for a span, applying profile overrides."""
        size = span.size
        if profile and not para.style.startswith("heading"):
            size = profile.get("body_size", size)
        elif profile and para.style.startswith("heading"):
            heading_sizes = profile.get("heading_sizes", {})
            size = heading_sizes.get(para.style, size)
        # Guard against zero/negative sizes
        if size <= 0:
            size = 10.0
        return size

    def _add_table(self, word_doc: Any, table: Any) -> None:
        """Add a formatted table to the Word document with proper borders,
        header styling, cell padding, and column widths."""
        from docx.shared import Pt, RGBColor, Emu
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        if table.num_rows == 0 or table.num_cols == 0:
            return

        doc_table = word_doc.add_table(
            rows=table.num_rows, cols=table.num_cols
        )
        doc_table.style = "Table Grid"

        # Disable autofit so we can control column widths
        doc_table.autofit = False

        # Calculate available page width (approximate — use 6.5 inches as default)
        try:
            section = word_doc.sections[-1]
            avail_width = section.page_width - section.left_margin - section.right_margin
        except Exception:
            avail_width = Emu(int(6.5 * 914400))  # 6.5 inches in EMU

        # Set uniform column widths
        if table.num_cols > 0:
            col_width = avail_width // table.num_cols
            for col_idx in range(table.num_cols):
                for row in doc_table.rows:
                    try:
                        row.cells[col_idx].width = col_width
                    except (IndexError, AttributeError):
                        pass

        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row):
                if c_idx >= table.num_cols or r_idx >= table.num_rows:
                    continue
                doc_cell = doc_table.cell(r_idx, c_idx)

                # Clear default paragraph and set text properly
                doc_cell.text = cell.text or ""

                # Apply explicit cell borders (ensures all cells are bordered
                # even if the style doesn't propagate)
                _set_cell_borders(doc_cell._element)

                # Set consistent cell padding
                _set_cell_margin(doc_cell._element, margin_twips=72)

                # Apply formatting to cell text
                for paragraph in doc_cell.paragraphs:
                    # Reduce paragraph spacing inside cells for compactness
                    paragraph.paragraph_format.space_before = Pt(1)
                    paragraph.paragraph_format.space_after = Pt(1)
                    for run in paragraph.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(9)
                        if cell.bold or cell.is_header:
                            run.font.bold = True

                # Header row shading — dark blue with white text
                if cell.is_header:
                    tc_pr = doc_cell._element.get_or_add_tcPr()
                    # Remove any existing shading first
                    existing_shd = tc_pr.find(qn("w:shd"))
                    if existing_shd is not None:
                        tc_pr.remove(existing_shd)
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), "1565C0")
                    shading.set(qn("w:val"), "clear")
                    tc_pr.append(shading)
                    for paragraph in doc_cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.bold = True

    def _add_image(self, word_doc: Any, para: Any) -> None:
        """Add an image to the Word document."""
        from docx.shared import Inches

        if not para.lines or not para.lines[0].spans:
            return

        span = para.lines[0].spans[0]
        data_uri = span.text

        if not data_uri.startswith("data:image/"):
            return

        try:
            import base64
            # Parse data URI
            header, b64_data = data_uri.split(",", 1)
            img_bytes = base64.b64decode(b64_data)

            # Get dimensions from span bbox
            width_pt = span.x1
            height_pt = span.y1

            # Write to temp buffer
            img_stream = io.BytesIO(img_bytes)

            # Scale to fit page width (max 6 inches)
            max_width = 6.0
            if width_pt > 0:
                width_inches = min(width_pt / 72.0, max_width)
            else:
                width_inches = 4.0

            word_doc.add_picture(img_stream, width=Inches(width_inches))
        except Exception as e:
            logger.debug(f"Failed to add image to DOCX: {e}")
