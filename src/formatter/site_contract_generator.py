"""
Site Contract Generator — fills a CTSA template with protocol-specific data.

Takes a blueprint CTSA template PDF and protocol extraction data,
produces a filled site contract maintaining the template's formatting.

Usage:
    generator = SiteContractGenerator()
    result = generator.generate(template_pdf_bytes, protocol_data)
    # result["docx_bytes"], result["html"], result["fill_report"]
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from typing import Any

from src.formatter.extractor import (
    FormattedDocument,
    FormattedPage,
    FormattedParagraph,
    FormattedSpan,
    FormattingExtractor,
)

logger = logging.getLogger(__name__)


@dataclass
class FillField:
    """A placeholder field in the template that needs filling."""
    placeholder: str       # e.g., "[add NHS TRUST NAME]"
    field_key: str         # e.g., "site_name"
    value: str = ""        # filled value
    page: int = 0
    filled: bool = False


@dataclass
class FillReport:
    """Report of what was filled in the template."""
    fields_found: int = 0
    fields_filled: int = 0
    fields_unfilled: int = 0
    details: list[dict[str, str]] = field(default_factory=list)


# Template placeholder patterns and their mappings
_PLACEHOLDER_MAPPINGS = [
    # Pattern in template → protocol data key
    (r"\[add NHS TRUST.*?\]", "site_name",
     "Site institution name"),
    (r"Dated:\s*_+", "agreement_date",
     "Agreement execution date"),
    (r"\(the \"Trial\"\)", "trial_name",
     "Trial name/title"),
    (r"\(the \"Protocol\"\)", "protocol_ref",
     "Protocol reference"),
    (r"\"Principal Investigator\"", "pi_name",
     "Principal Investigator name"),
    (r"\[insert.*?\]", "generic_insert",
     "Generic insertion point"),
]


class SiteContractGenerator:
    """Generates a filled site contract from a template + protocol data."""

    def __init__(self):
        self.extractor = FormattingExtractor()

    def generate(
        self,
        template_pdf_bytes: bytes,
        protocol_data: dict[str, Any],
        site_info: dict[str, str] | None = None,
    ) -> dict[str, Any]:
        """Generate a filled site contract.

        Args:
            template_pdf_bytes: The CTSA template PDF
            protocol_data: Protocol extraction data (from pipeline output)
            site_info: Optional site-specific information to fill

        Returns:
            dict with "docx_bytes", "html", "fill_report"
        """
        # Extract template formatting
        template_doc = self.extractor.extract(template_pdf_bytes, "template")

        # Build fill values from protocol data
        fill_values = self._build_fill_values(protocol_data, site_info)

        # Generate filled HTML with template formatting
        html = self._render_filled_html(template_doc, fill_values)

        # Generate filled DOCX
        docx_bytes = self._render_filled_docx(template_doc, fill_values)

        # Build fill report
        fill_report = self._build_fill_report(template_doc, fill_values)

        return {
            "html": html,
            "docx_bytes": docx_bytes,
            "fill_report": fill_report,
            "template_pages": len(template_doc.pages),
            "template_paragraphs": template_doc.total_paragraphs,
        }

    def _build_fill_values(
        self,
        protocol_data: dict[str, Any],
        site_info: dict[str, str] | None = None,
    ) -> dict[str, str]:
        """Extract fill values from protocol data and site info."""
        meta = protocol_data.get("metadata", {})
        site = site_info or {}

        # Protocol-derived values
        sponsor = meta.get("sponsor", "")
        title = meta.get("title", "")
        protocol_num = meta.get("protocol_number", "")
        phase = meta.get("phase", "")
        indication = meta.get("indication", "")
        ta = meta.get("therapeutic_area", "")

        # Extract procedures for Appendix
        procedures = []
        for table in protocol_data.get("tables", []):
            for proc in table.get("procedures", []):
                name = proc.get("canonical_name", proc.get("raw_name", ""))
                cpt = proc.get("cpt_code", proc.get("code", ""))
                if name and name not in [p["name"] for p in procedures]:
                    procedures.append({"name": name, "cpt": cpt})

        # Build visit schedule summary
        visits = []
        for table in protocol_data.get("tables", []):
            for vw in table.get("visit_windows", []):
                name = vw.get("visit_name", vw.get("name", ""))
                if name and name not in visits:
                    visits.append(name)

        return {
            # Header fields
            "sponsor_name": sponsor,
            "trial_title": title,
            "protocol_number": protocol_num,
            "phase": phase,
            "indication": indication,
            "therapeutic_area": ta,

            # Site fields
            "site_name": site.get("name", "[Site Institution Name]"),
            "site_address": site.get("address", "[Site Address]"),
            "pi_name": site.get("pi_name", "[Principal Investigator Name]"),
            "pi_qualifications": site.get("pi_qualifications", "[PI Qualifications]"),

            # Date fields
            "agreement_date": site.get("date", "[Date]"),
            "effective_date": site.get("effective_date", "[Effective Date]"),

            # Protocol-derived content
            "procedures_list": procedures,
            "visits_list": visits,
            "num_procedures": str(len(procedures)),
            "num_visits": str(len(visits)),

            # Trial drug
            "trial_drug": self._extract_trial_drug(protocol_data),
        }

    def _extract_trial_drug(self, protocol_data: dict[str, Any]) -> str:
        """Extract trial drug name from protocol data."""
        meta = protocol_data.get("metadata", {})
        title = meta.get("title", "")

        # Common patterns in protocol titles
        drug_patterns = [
            r"(?:study of|evaluate|evaluating)\s+(\w+(?:\s+\w+)?)",
            r"(\w+(?:-\w+)?)\s+(?:versus|vs\.?|compared)",
        ]
        for pattern in drug_patterns:
            match = re.search(pattern, title, re.IGNORECASE)
            if match:
                return match.group(1)

        return "[Trial Drug Name]"

    def _render_filled_html(
        self,
        template_doc: FormattedDocument,
        fill_values: dict[str, Any],
    ) -> str:
        """Render the filled template as HTML."""
        parts = []
        parts.append("""<!DOCTYPE html>
<html>
<head>
<style>
  body { font-family: Arial, sans-serif; font-size: 11pt; line-height: 1.4;
         max-width: 800px; margin: 40px auto; color: #333; }
  h1 { font-size: 14pt; font-weight: bold; text-align: center; margin: 20px 0; }
  h2 { font-size: 12pt; font-weight: bold; margin: 16px 0 8px; }
  h3 { font-size: 11pt; font-weight: bold; margin: 12px 0 6px; }
  .filled { background-color: #e8f5e9; padding: 1px 4px; border-radius: 2px;
            border-bottom: 2px solid #4caf50; }
  .unfilled { background-color: #fff3e0; padding: 1px 4px; border-radius: 2px;
              border-bottom: 2px solid #ff9800; color: #e65100; }
  .section-num { color: #1565c0; font-weight: bold; margin-right: 8px; }
  .page-break { border-top: 1px dashed #ccc; margin: 30px 0; padding-top: 10px; }
  .appendix-header { font-size: 13pt; font-weight: bold; text-align: center;
                     margin: 30px 0 15px; text-decoration: underline; }
  table.procedures { border-collapse: collapse; width: 100%; margin: 12px 0; }
  table.procedures th, table.procedures td {
    border: 1px solid #ccc; padding: 6px 10px; text-align: left; font-size: 10pt; }
  table.procedures th { background: #1565c0; color: white; font-weight: bold; }
  table.procedures tr:nth-child(even) { background: #f5f5f5; }
  .header-bar { background: #f5f5f5; padding: 8px 16px; border-bottom: 2px solid #1565c0;
                font-size: 9pt; color: #666; margin-bottom: 20px; }
  .signature-line { border-bottom: 1px solid #333; width: 300px; display: inline-block;
                    margin: 8px 0; }
  .meta-table { width: 100%; margin: 16px 0; }
  .meta-table td { padding: 4px 8px; vertical-align: top; }
  .meta-table .label { font-weight: bold; width: 180px; color: #555; }
</style>
</head>
<body>
""")

        # Protocol info header
        parts.append(f"""
<div class="header-bar">
  <strong>Clinical Trial Site Agreement</strong> |
  Protocol: <span class="filled">{fill_values['protocol_number']}</span> |
  Sponsor: <span class="filled">{fill_values['sponsor_name']}</span> |
  Phase: <span class="filled">{fill_values['phase']}</span>
</div>
""")

        # Title page
        parts.append('<h1>Clinical Trial Site Agreement</h1>')
        parts.append(f"""
<p style="text-align:center; margin: 20px 0;">
  Dated: <span class="{'filled' if fill_values['agreement_date'] != '[Date]' else 'unfilled'}">{fill_values['agreement_date']}</span>
</p>

<p><strong>Between:</strong></p>

<p>(1) <span class="filled">{fill_values['sponsor_name']}</span>
(the "<strong>Sponsor</strong>");</p>

<p>and</p>

<p>(2) <span class="{'filled' if fill_values['site_name'] != '[Site Institution Name]' else 'unfilled'}">{fill_values['site_name']}</span>
(the "<strong>Site</strong>").</p>

<p>together: "<strong>the Parties</strong>"</p>
""")

        # Recitals
        parts.append(f"""
<h2>Whereas:</h2>

<p>(A) The Sponsor is coordinating the trial entitled:
<span class="filled">{fill_values['trial_title'][:120]}</span>
(the "<strong>Trial</strong>"), Protocol Number:
<span class="filled">{fill_values['protocol_number']}</span>,
which will be conducted according to the Protocol.</p>

<p>(B) The Trial is a <span class="filled">{fill_values['phase']}</span> study
in the <span class="filled">{fill_values['therapeutic_area']}</span> therapeutic area,
investigating <span class="filled">{fill_values['indication']}</span>.</p>

<p>(C) The Sponsor wishes to appoint the Site to participate in the Trial and the Site
has agreed to participate on the terms set out in this Agreement.</p>

<p>(D) The Principal Investigator,
<span class="{'filled' if fill_values['pi_name'] != '[Principal Investigator Name]' else 'unfilled'}">{fill_values['pi_name']}</span>,
has agreed to conduct the Trial at the Site in accordance with the Protocol.</p>

<p>(E) The Sponsor has entered into an agreement with the Supplier to cover the supply
of the Trial Drug (<span class="filled">{fill_values['trial_drug']}</span>).</p>
""")

        # Render template content with formatting preserved
        for page in template_doc.pages:
            # Skip first page (we replaced it with our filled version)
            if page.page_number == 0:
                continue

            # Skip header/footer content
            for para in page.paragraphs:
                text = para.text.strip()
                # Skip template header/footer lines
                if any(skip in text for skip in [
                    "UCL CTC CTSA template",
                    "Modified for ANIMATE",
                    "Page ", "of 31",
                ]):
                    continue
                if not text:
                    continue

                # Determine HTML tag
                tag = "p"
                extra_class = ""
                if para.style.startswith("heading"):
                    level = para.style[-1] if para.style[-1].isdigit() else "3"
                    tag = f"h{level}"
                elif "APPENDIX" in text.upper() and len(text) < 50:
                    tag = "div"
                    extra_class = ' class="appendix-header"'

                # Build span HTML
                spans_html = []
                for line in para.lines:
                    for span in line.spans:
                        t = span.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

                        # Fill placeholders
                        t = self._fill_placeholders(t, fill_values)

                        if span.bold:
                            t = f"<strong>{t}</strong>"
                        if span.italic:
                            t = f"<em>{t}</em>"
                        if span.underline:
                            t = f"<u>{t}</u>"
                        if span.superscript:
                            t = f"<sup>{t}</sup>"
                        if span.subscript:
                            t = f"<sub>{t}</sub>"
                        spans_html.append(t)

                content = "".join(spans_html)
                parts.append(f"<{tag}{extra_class}>{content}</{tag}>")

        # Appendix: Protocol Summary
        parts.append("""
<div class="page-break"></div>
<div class="appendix-header">Protocol Summary (Auto-Generated from Extraction)</div>
""")

        parts.append(f"""
<table class="meta-table">
  <tr><td class="label">Protocol Number:</td><td><span class="filled">{fill_values['protocol_number']}</span></td></tr>
  <tr><td class="label">Protocol Title:</td><td><span class="filled">{fill_values['trial_title'][:150]}</span></td></tr>
  <tr><td class="label">Sponsor:</td><td><span class="filled">{fill_values['sponsor_name']}</span></td></tr>
  <tr><td class="label">Phase:</td><td><span class="filled">{fill_values['phase']}</span></td></tr>
  <tr><td class="label">Therapeutic Area:</td><td><span class="filled">{fill_values['therapeutic_area']}</span></td></tr>
  <tr><td class="label">Indication:</td><td><span class="filled">{fill_values['indication']}</span></td></tr>
  <tr><td class="label">Trial Drug:</td><td><span class="filled">{fill_values['trial_drug']}</span></td></tr>
  <tr><td class="label">Total Procedures:</td><td>{fill_values['num_procedures']}</td></tr>
  <tr><td class="label">Total Visits:</td><td>{fill_values['num_visits']}</td></tr>
</table>
""")

        # Procedures table
        procedures = fill_values.get("procedures_list", [])
        if procedures:
            parts.append('<h3>Schedule of Procedures</h3>')
            parts.append('<table class="procedures">')
            parts.append('<thead><tr><th>#</th><th>Procedure</th><th>CPT Code</th></tr></thead>')
            parts.append('<tbody>')
            for i, proc in enumerate(procedures, 1):
                parts.append(f'<tr><td>{i}</td><td>{proc["name"]}</td><td>{proc.get("cpt", "—")}</td></tr>')
            parts.append('</tbody></table>')

        # Visits
        visits = fill_values.get("visits_list", [])
        if visits:
            parts.append('<h3>Study Visits</h3>')
            parts.append('<table class="procedures">')
            parts.append('<thead><tr><th>#</th><th>Visit</th></tr></thead>')
            parts.append('<tbody>')
            for i, visit in enumerate(visits, 1):
                parts.append(f'<tr><td>{i}</td><td>{visit}</td></tr>')
            parts.append('</tbody></table>')

        # Signature block
        parts.append("""
<div class="page-break"></div>
<h2>Signatures</h2>

<p>IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date first written above.</p>

<table class="meta-table" style="margin-top: 30px;">
  <tr>
    <td style="width:50%">
      <p><strong>For and on behalf of the Sponsor:</strong></p>
      <p><br/><span class="signature-line">&nbsp;</span></p>
      <p>Name: <span class="signature-line">&nbsp;</span></p>
      <p>Title: <span class="signature-line">&nbsp;</span></p>
      <p>Date: <span class="signature-line">&nbsp;</span></p>
    </td>
    <td style="width:50%">
      <p><strong>For and on behalf of the Site:</strong></p>
      <p><br/><span class="signature-line">&nbsp;</span></p>
      <p>Name: <span class="signature-line">&nbsp;</span></p>
      <p>Title: <span class="signature-line">&nbsp;</span></p>
      <p>Date: <span class="signature-line">&nbsp;</span></p>
    </td>
  </tr>
</table>
""")

        parts.append("</body></html>")
        return "\n".join(parts)

    def _fill_placeholders(self, text: str, fill_values: dict[str, Any]) -> str:
        """Replace placeholder text with filled values."""
        replacements = [
            (r"\[add NHS TRUST.*?\]", fill_values.get("site_name", "[Site Name]")),
            (r"\[insert.*?name.*?\]", fill_values.get("pi_name", "[PI Name]")),
            (r"\[insert.*?address.*?\]", fill_values.get("site_address", "[Address]")),
            (r"the \"Trial\"", f'the Trial (<span class="filled">{fill_values.get("protocol_number", "")}</span>)'),
            (r"ANIMATE", fill_values.get("protocol_number", "ANIMATE")),
        ]

        result = text
        for pattern, replacement in replacements:
            result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
        return result

    def _render_filled_docx(
        self,
        template_doc: FormattedDocument,
        fill_values: dict[str, Any],
    ) -> bytes:
        """Render the filled template as a DOCX document."""
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Set A4 page size and margins
        for section in doc.sections:
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Clinical Trial Site Agreement")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = "Arial"

        # Parties
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Dated: {fill_values['agreement_date']}")
        run.font.size = Pt(11)
        run.font.name = "Arial"

        doc.add_paragraph()  # blank line

        p = doc.add_paragraph()
        run = p.add_run("Between:")
        run.font.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(11)

        p = doc.add_paragraph()
        run = p.add_run(f"(1) {fill_values['sponsor_name']} ")
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run = p.add_run('(the "Sponsor");')
        run.font.name = "Arial"
        run.font.size = Pt(11)

        p = doc.add_paragraph("and")
        p.runs[0].font.name = "Arial"

        p = doc.add_paragraph()
        run = p.add_run(f"(2) {fill_values['site_name']} ")
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run = p.add_run('(the "Site").')
        run.font.name = "Arial"
        run.font.size = Pt(11)

        p = doc.add_paragraph()
        run = p.add_run('together: "the Parties"')
        run.font.name = "Arial"
        run.font.size = Pt(11)

        # Recitals
        doc.add_paragraph()
        p = doc.add_heading("Whereas:", level=2)

        recitals = [
            f"(A) The Sponsor is coordinating the trial entitled: {fill_values['trial_title'][:120]} "
            f"(the \"Trial\"), Protocol Number: {fill_values['protocol_number']}.",
            f"(B) The Trial is a {fill_values['phase']} study in the {fill_values['therapeutic_area']} "
            f"therapeutic area, investigating {fill_values['indication']}.",
            "(C) The Sponsor wishes to appoint the Site to participate in the Trial and the Site "
            "has agreed to participate on the terms set out in this Agreement.",
            f"(D) The Principal Investigator, {fill_values['pi_name']}, has agreed to conduct "
            "the Trial at the Site in accordance with the Protocol.",
            f"(E) The Sponsor has entered into an agreement for the supply of the Trial Drug "
            f"({fill_values['trial_drug']}).",
        ]
        for recital in recitals:
            p = doc.add_paragraph(recital)
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)

        # Render template body content
        for page in template_doc.pages:
            if page.page_number < 4:  # skip pages we replaced
                continue

            for para in page.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                # Skip headers/footers
                if any(skip in text for skip in [
                    "UCL CTC CTSA template", "Modified for ANIMATE", "Page ", "of 31",
                ]):
                    continue

                if para.style.startswith("heading"):
                    level = int(para.style[-1]) if para.style[-1].isdigit() else 3
                    level = min(level, 4)
                    p = doc.add_heading(text, level=level)
                else:
                    # Fill placeholders
                    filled_text = text
                    filled_text = re.sub(r"ANIMATE", fill_values.get("protocol_number", "ANIMATE"), filled_text)
                    filled_text = re.sub(r"\[add NHS TRUST.*?\]", fill_values.get("site_name", "[Site]"), filled_text, flags=re.IGNORECASE)

                    p = doc.add_paragraph(filled_text)
                    for run in p.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(11)

                    if para.is_bold:
                        for run in p.runs:
                            run.font.bold = True

        # Protocol Summary appendix
        doc.add_page_break()
        doc.add_heading("Protocol Summary (Auto-Generated)", level=1)

        # Summary table
        meta_fields = [
            ("Protocol Number", fill_values['protocol_number']),
            ("Sponsor", fill_values['sponsor_name']),
            ("Phase", fill_values['phase']),
            ("Therapeutic Area", fill_values['therapeutic_area']),
            ("Indication", fill_values['indication']),
            ("Trial Drug", fill_values['trial_drug']),
            ("Total Procedures", fill_values['num_procedures']),
            ("Total Visits", fill_values['num_visits']),
        ]
        table = doc.add_table(rows=len(meta_fields), cols=2)
        table.style = "Table Grid"
        for i, (label, value) in enumerate(meta_fields):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            for cell in [table.cell(i, 0), table.cell(i, 1)]:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(10)
            # Bold labels
            for run in table.cell(i, 0).paragraphs[0].runs:
                run.font.bold = True

        # Procedures table
        procedures = fill_values.get("procedures_list", [])
        if procedures:
            doc.add_paragraph()
            doc.add_heading("Schedule of Procedures", level=2)
            proc_table = doc.add_table(rows=len(procedures) + 1, cols=3)
            proc_table.style = "Table Grid"
            headers = ["#", "Procedure", "CPT Code"]
            for j, h in enumerate(headers):
                proc_table.cell(0, j).text = h
                for run in proc_table.cell(0, j).paragraphs[0].runs:
                    run.font.bold = True
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
            for i, proc in enumerate(procedures, 1):
                proc_table.cell(i, 0).text = str(i)
                proc_table.cell(i, 1).text = proc["name"]
                proc_table.cell(i, 2).text = proc.get("cpt") or "—"

        # Signature page
        doc.add_page_break()
        doc.add_heading("Signatures", level=1)
        p = doc.add_paragraph(
            "IN WITNESS WHEREOF, the Parties have executed this Agreement "
            "as of the date first written above."
        )
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)

        for party in ["Sponsor", "Site"]:
            doc.add_paragraph()
            p = doc.add_paragraph(f"For and on behalf of the {party}:")
            for run in p.runs:
                run.font.bold = True
                run.font.name = "Arial"
            for label in ["Signature: ___________________________",
                         "Name: ___________________________",
                         "Title: ___________________________",
                         "Date: ___________________________"]:
                p = doc.add_paragraph(label)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(11)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _build_fill_report(
        self,
        template_doc: FormattedDocument,
        fill_values: dict[str, Any],
    ) -> dict[str, Any]:
        """Build a report of what was filled in the template."""
        filled = []
        unfilled = []

        field_checks = [
            ("sponsor_name", "Sponsor Name"),
            ("trial_title", "Trial Title"),
            ("protocol_number", "Protocol Number"),
            ("phase", "Phase"),
            ("indication", "Indication"),
            ("therapeutic_area", "Therapeutic Area"),
            ("trial_drug", "Trial Drug"),
            ("site_name", "Site Name"),
            ("pi_name", "Principal Investigator"),
            ("agreement_date", "Agreement Date"),
        ]

        for key, label in field_checks:
            val = fill_values.get(key, "")
            if val and not val.startswith("["):
                filled.append({"field": label, "value": str(val)[:80]})
            else:
                unfilled.append({"field": label, "placeholder": str(val)})

        return {
            "fields_filled": len(filled),
            "fields_unfilled": len(unfilled),
            "filled": filled,
            "unfilled": unfilled,
            "procedures_count": len(fill_values.get("procedures_list", [])),
            "visits_count": len(fill_values.get("visits_list", [])),
        }
