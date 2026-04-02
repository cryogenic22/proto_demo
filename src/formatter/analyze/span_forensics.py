"""
Span-Level Forensic Diff — measures exactly where formatting is lost.

Compares three levels:
1. Source (raw PyMuPDF spans) — ground truth
2. IR (FormattedDocument after extraction) — what the ingestor captured
3. Output (re-ingested rendered document) — what survived rendering

For each level, counts formatting attributes (bold, italic, underline,
superscript, subscript, colored) and reports which specific spans dropped
and WHY (font-flag misread, synthetic bold, inherited style, renderer bug).

Usage:
    forensics = SpanForensics()
    report = forensics.analyze_pdf(pdf_bytes)
    print(report.summary())
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from typing import Any

import fitz

from src.formatter.extractor import FormattingExtractor, FormattedDocument

logger = logging.getLogger(__name__)


@dataclass
class SpanStats:
    """Counts of formatting attributes at a given level."""
    total_spans: int = 0
    bold: int = 0
    italic: int = 0
    underline: int = 0
    superscript: int = 0
    subscript: int = 0
    colored: int = 0  # non-black
    fonts: dict[str, int] = field(default_factory=dict)
    sizes: dict[float, int] = field(default_factory=dict)


@dataclass
class DroppedSpan:
    """A specific span where formatting was lost."""
    text: str
    page: int
    attribute: str  # bold, italic, underline, etc.
    root_cause: str  # font_flag_misread, synthetic_bold, inherited_style, renderer_bug
    source_value: str = ""
    ir_value: str = ""


@dataclass
class ForensicReport:
    """Complete forensic analysis of formatting fidelity."""
    source_stats: SpanStats = field(default_factory=SpanStats)
    ir_stats: SpanStats = field(default_factory=SpanStats)
    output_stats: SpanStats | None = None
    dropped_spans: list[DroppedSpan] = field(default_factory=list)
    root_cause_counts: dict[str, int] = field(default_factory=dict)

    def fidelity(self, attribute: str) -> float:
        """Compute fidelity ratio for an attribute (IR / source)."""
        src = getattr(self.source_stats, attribute, 0)
        ir = getattr(self.ir_stats, attribute, 0)
        if src == 0:
            return 1.0 if ir == 0 else 0.0
        return min(ir / src, 1.0)

    def render_fidelity(self, attribute: str) -> float:
        """Compute fidelity ratio for an attribute (output / source)."""
        if not self.output_stats:
            return 0.0
        src = getattr(self.source_stats, attribute, 0)
        out = getattr(self.output_stats, attribute, 0)
        if src == 0:
            return 1.0 if out == 0 else 0.0
        return min(out / src, 1.0)

    def summary(self) -> str:
        """Human-readable summary."""
        lines = ["FORENSIC SPAN ANALYSIS", "=" * 50, ""]

        lines.append("ATTRIBUTE COUNTS:")
        lines.append(f"{'Attribute':<15} {'Source':>8} {'IR':>8} {'Fidelity':>10}")
        lines.append("-" * 45)
        for attr in ["bold", "italic", "underline", "superscript", "subscript", "colored"]:
            src = getattr(self.source_stats, attr, 0)
            ir = getattr(self.ir_stats, attr, 0)
            fid = self.fidelity(attr)
            lines.append(f"{attr:<15} {src:>8} {ir:>8} {fid:>9.0%}")

        if self.output_stats:
            lines.append("")
            lines.append("RENDER PIPELINE (Source → IR → Output):")
            lines.append(f"{'Attribute':<15} {'Source':>8} {'IR':>8} {'Output':>8} {'End-to-End':>10}")
            lines.append("-" * 55)
            for attr in ["bold", "italic", "underline", "superscript", "subscript", "colored"]:
                src = getattr(self.source_stats, attr, 0)
                ir = getattr(self.ir_stats, attr, 0)
                out = getattr(self.output_stats, attr, 0)
                e2e = self.render_fidelity(attr)
                lines.append(f"{attr:<15} {src:>8} {ir:>8} {out:>8} {e2e:>9.0%}")

        if self.root_cause_counts:
            lines.append("")
            lines.append("ROOT CAUSES OF DROPPED SPANS:")
            for cause, count in sorted(self.root_cause_counts.items(), key=lambda x: -x[1]):
                lines.append(f"  {cause}: {count}")

        if self.dropped_spans:
            lines.append("")
            lines.append(f"SAMPLE DROPPED SPANS (first 10 of {len(self.dropped_spans)}):")
            for ds in self.dropped_spans[:10]:
                lines.append(f"  p{ds.page} [{ds.attribute}] \"{ds.text[:40]}\" — {ds.root_cause}")

        return "\n".join(lines)


# Bold detection patterns in font names
_BOLD_NAME_PATTERNS = re.compile(
    r"Bold|Bd|Demi|Heavy|Black|Semibold|ExtraBold|UltraBold",
    re.IGNORECASE,
)

_ITALIC_NAME_PATTERNS = re.compile(
    r"Italic|It|Oblique|Slanted|Inclined",
    re.IGNORECASE,
)


class SpanForensics:
    """Forensic analysis of formatting fidelity at the span level."""

    def __init__(self):
        self.extractor = FormattingExtractor()

    def analyze_pdf(self, pdf_bytes: bytes, filename: str = "") -> ForensicReport:
        """Full forensic analysis: source → IR → DOCX output."""
        report = ForensicReport()

        # Step 1: Extract ground truth from raw PyMuPDF
        report.source_stats = self._extract_source_stats(pdf_bytes)

        # Step 2: Extract via our IR
        ir = self.extractor.extract(pdf_bytes, filename)
        report.ir_stats = self._extract_ir_stats(ir)

        # Step 3: Identify dropped spans and root causes
        self._find_dropped_spans(pdf_bytes, ir, report)

        # Step 4: Render to DOCX and re-analyze
        try:
            from src.formatter.docx_renderer import DOCXRenderer
            renderer = DOCXRenderer()
            docx_bytes = renderer.render(ir)
            report.output_stats = self._extract_docx_stats(docx_bytes)
        except Exception as e:
            logger.warning(f"DOCX render/re-analysis failed: {e}")

        return report

    def _extract_source_stats(self, pdf_bytes: bytes) -> SpanStats:
        """Extract raw span counts directly from PyMuPDF (ground truth)."""
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        stats = SpanStats()

        for page_idx in range(doc.page_count):
            page = doc[page_idx]
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text:
                            continue

                        stats.total_spans += 1
                        flags = span.get("flags", 0)
                        font_name = span.get("font", "")
                        color = span.get("color", 0)

                        # Bold: flag bit 4 (16) OR font name contains Bold
                        is_bold = bool(flags & 16) or bool(_BOLD_NAME_PATTERNS.search(font_name))
                        if is_bold:
                            stats.bold += 1

                        # Italic: flag bit 1 (2) OR font name contains Italic
                        is_italic = bool(flags & 2) or bool(_ITALIC_NAME_PATTERNS.search(font_name))
                        if is_italic:
                            stats.italic += 1

                        # Underline: flag bit 2 (4)
                        if flags & 4:
                            stats.underline += 1

                        # Superscript: flag bit 0 (1)
                        if flags & 1:
                            stats.superscript += 1

                        # Subscript: small font (heuristic)
                        size = span.get("size", 10.0)
                        if size < 7.5 and not (flags & 1):
                            stats.subscript += 1

                        # Colored (non-black)
                        if color != 0:
                            r = (color >> 16) & 0xFF
                            g = (color >> 8) & 0xFF
                            b = color & 0xFF
                            if r > 20 or g > 20 or b > 20:
                                stats.colored += 1

                        # Font tracking
                        base_font = font_name.split("-")[0].split(",")[0].replace("MT", "").replace("PS", "").strip()
                        stats.fonts[base_font] = stats.fonts.get(base_font, 0) + 1

                        # Size tracking
                        rounded_size = round(size, 1)
                        stats.sizes[rounded_size] = stats.sizes.get(rounded_size, 0) + 1

        doc.close()
        return stats

    def _extract_ir_stats(self, doc: FormattedDocument) -> SpanStats:
        """Extract span counts from our IR."""
        stats = SpanStats()

        for page in doc.pages:
            for para in page.paragraphs:
                for line in para.lines:
                    for span in line.spans:
                        if not span.text.strip():
                            continue

                        stats.total_spans += 1

                        if span.bold:
                            stats.bold += 1
                        if span.italic:
                            stats.italic += 1
                        if span.underline:
                            stats.underline += 1
                        if span.superscript:
                            stats.superscript += 1
                        if span.subscript:
                            stats.subscript += 1

                        r, g, b = span.color_rgb
                        if r > 20 or g > 20 or b > 20:
                            stats.colored += 1

                        base_font = span.font_family
                        stats.fonts[base_font] = stats.fonts.get(base_font, 0) + 1

                        rounded_size = round(span.size, 1)
                        stats.sizes[rounded_size] = stats.sizes.get(rounded_size, 0) + 1

        return stats

    def _extract_docx_stats(self, docx_bytes: bytes) -> SpanStats:
        """Extract span counts from a rendered DOCX."""
        import io
        from docx import Document
        from docx.shared import RGBColor

        stats = SpanStats()
        word_doc = Document(io.BytesIO(docx_bytes))

        for para in word_doc.paragraphs:
            for run in para.runs:
                text = run.text.strip()
                if not text:
                    continue

                stats.total_spans += 1

                if run.font.bold:
                    stats.bold += 1
                if run.font.italic:
                    stats.italic += 1
                if run.font.underline:
                    stats.underline += 1
                if run.font.superscript:
                    stats.superscript += 1
                if run.font.subscript:
                    stats.subscript += 1

                try:
                    if run.font.color and run.font.color.rgb and run.font.color.rgb != RGBColor(0, 0, 0):
                        stats.colored += 1
                except Exception:
                    pass

                if run.font.name:
                    stats.fonts[run.font.name] = stats.fonts.get(run.font.name, 0) + 1

                if run.font.size:
                    size = round(run.font.size.pt, 1)
                    stats.sizes[size] = stats.sizes.get(size, 0) + 1

        return stats

    def _find_dropped_spans(
        self, pdf_bytes: bytes, ir: FormattedDocument, report: ForensicReport,
    ) -> None:
        """Identify specific dropped spans and categorize root causes."""
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")

        for page_idx in range(doc.page_count):
            page = doc[page_idx]
            blocks = page.get_text("dict")["blocks"]

            for block in blocks:
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text or len(text) < 2:
                            continue

                        flags = span.get("flags", 0)
                        font_name = span.get("font", "")

                        # Check: is this span bold in source but not in IR?
                        src_bold = bool(flags & 16) or bool(_BOLD_NAME_PATTERNS.search(font_name))
                        ir_bold = self._find_span_in_ir(ir, page_idx, text, "bold")

                        if src_bold and not ir_bold:
                            # Determine root cause
                            if not (flags & 16) and _BOLD_NAME_PATTERNS.search(font_name):
                                cause = "font_name_bold_not_flag"
                            elif flags & 16:
                                cause = "flag_present_but_ir_missed"
                            else:
                                cause = "unknown_bold_drop"

                            report.dropped_spans.append(DroppedSpan(
                                text=text[:50], page=page_idx + 1,
                                attribute="bold", root_cause=cause,
                                source_value=f"flags={flags}, font={font_name}",
                            ))
                            report.root_cause_counts[cause] = report.root_cause_counts.get(cause, 0) + 1

                        # Check italic
                        src_italic = bool(flags & 2) or bool(_ITALIC_NAME_PATTERNS.search(font_name))
                        ir_italic = self._find_span_in_ir(ir, page_idx, text, "italic")

                        if src_italic and not ir_italic:
                            if not (flags & 2) and _ITALIC_NAME_PATTERNS.search(font_name):
                                cause = "font_name_italic_not_flag"
                            elif flags & 2:
                                cause = "flag_present_but_ir_missed"
                            else:
                                cause = "unknown_italic_drop"

                            report.dropped_spans.append(DroppedSpan(
                                text=text[:50], page=page_idx + 1,
                                attribute="italic", root_cause=cause,
                                source_value=f"flags={flags}, font={font_name}",
                            ))
                            report.root_cause_counts[cause] = report.root_cause_counts.get(cause, 0) + 1

        doc.close()

    def _find_span_in_ir(
        self, ir: FormattedDocument, page_idx: int, text: str, attribute: str,
    ) -> bool:
        """Check if a span with given text has the attribute set in the IR."""
        if page_idx >= len(ir.pages):
            return False

        page = ir.pages[page_idx]
        text_lower = text.lower()[:20]

        for para in page.paragraphs:
            for line in para.lines:
                for span in line.spans:
                    if text_lower in span.text.lower():
                        return getattr(span, attribute, False)

        return False
